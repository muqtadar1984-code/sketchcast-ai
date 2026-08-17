"""Coverage gate — did the generated artifact actually teach what the analyzer
said this chapter contains?

Agent 2 publishes, per episode, the sections it chunked (``sections_covered``)
and the concept ids it extracted (``key_concepts_introduced``). Agent 3's prompt
is built from exactly those two lists — ``_build_episode_context`` writes them
into "Sections covered:" and "KEY CONCEPTS TO TEACH". Until now NOTHING compared
the reply back to them, so a script that silently taught a third of the chapter
finished as a clean success: full progress bar, video rendered, credit consumed.

That was tolerable while every call ran on Sonnet. It stops being tolerable the
moment a cheaper model is introduced, because THINNING is how a cheaper model
fails and it fails silently — no exception, no truncation, just less teaching.
The pipeline has already proved it will report success on garbage: on 2026-08-09
a teacher generated 17 kits, zero errors, from a text layer that turned out to be
unreadable mojibake.

So this module measures one number per generated artifact and records it. The
number is the point: with it on every row, "is Haiku thinner than Sonnet?" is a
SQL query over ``generations.params.coverage`` instead of an argument. Without
it, the flip is unfalsifiable.

Pure and cheap by construction: NO model call, no network, no I/O except the
optional .docx reader at the bottom. It runs on every generation, so it is set
arithmetic over a few thousand tokens (measured ~4 ms for a 14-segment script).
"""

from __future__ import annotations

import os
import unicodedata

from shared.part_label import names_nothing
from shared.scripts import SPACELESS_SCRIPTS, script_of


# ── What counts as a topic ───────────────────────────────────────────────────
# Two units, both already carried by the pipeline, pooled into one list:
#
#   SECTIONS  — episode["sections_covered"], the structural headings the chunker
#               put in this part. The founder's unit, and the honest one… when
#               the book has headings. It very often doesn't: the structurer
#               emits the literal placeholder "Content" when a chapter's page
#               range yields no headings, and that placeholder was 143 of the
#               304 parts in production (47%). Measuring a script against the
#               word "Content" is not a coverage test, so headings that name
#               nothing are dropped by the SAME rule that stops them being
#               displayed — shared/part_label.names_nothing, not a second copy
#               of it. On the one real analysis+script pair in the repo,
#               sections_covered is ["Content"] and contributes nothing.
#
#   CONCEPTS  — the analyzer's own extracted concept names, scoped to this
#               episode by key_concepts_introduced. This is the segment-to-
#               section mapping the pipeline already carries, and it is the
#               sharper detector: concept names are distinctive noun phrases
#               ("Broadcast Message", "Data Logger"), they are what the prompt
#               explicitly ORDERS the script to teach, and there are ~20 of them
#               per chapter where there may be one usable section heading.
#
# Pooling them means the gate still has a denominator on the 47% of parts whose
# headings are placeholders, and still honours real headings where they exist.


# ── How "addressed" is measured ──────────────────────────────────────────────
# Not exact heading match: generated prose rightly paraphrases, and a script
# that teaches "the light reactions" has not failed to teach "light-dependent
# reactions". Not a length heuristic either: a cheap model's failure is fluent
# and full-length, it is just about less. What survives both objections is
# per-topic LEXICAL RECALL — does the artifact's text contain the words that
# make this topic that topic and not its neighbour?
#
# 6 characters. Long enough to keep distinct technical terms apart in the sample
# chapter measured ("interaction"→intera vs "interval"→interv, "statistical" and
# "statistics"→statis), short enough to collapse the inflection that would
# otherwise read as a miss ("photosynthesis"/"photosynthetic"→photos,
# "designer"/"designed"/"design"→design). Tokens shorter than this are matched
# EXACTLY instead — a 3-letter name like "Air" must not match "chair".
_STEM = 6

# A topic counts as addressed when at least half its distinguishing tokens
# appear. Half, not all: the model paraphrases, and requiring every token flags
# good scripts for using the teacher's wording instead of the analyzer's — on
# the real script measured below, all-tokens and half-tokens agree on the whole
# script but diverge (0.826 vs 0.870) the moment it degrades, and the stricter
# rule is the one that moves for the wrong reason. Not "any one token" either:
# most concept names in one chapter share a head word ("cell membrane", "cell
# wall", "cell division"), so a single-token rule scores a script that only ever
# says "cell" as covering all three. The distinctiveness filter below is what
# makes "half" safe.
_ADDRESSED_SHARE = 0.5

# A token that appears in two or more of this episode's topic names cannot tell
# those topics apart, so it is dropped before the share above is computed:
# "Data Logger" / "Data Collection Interval" / "Data Representation" are
# measured on logger / collection+interval / representation, not on "data",
# which all three would match from a single mention. A topic whose name is
# ENTIRELY made of shared words keeps its raw tokens rather than becoming
# unmeasurable — a weak key still beats no key.
_DROP_TOKEN_AT_DF = 2

# …and at most this many keys per topic. A heading longer than eight content
# words is a sentence, and asking whether half a sentence's words appear is a
# paraphrase test, not a coverage test. Keys are taken rarest-first, so the
# cap keeps the discriminating words and drops the filler.
_MAX_TOPIC_KEYS = 8


# ── Thresholds ───────────────────────────────────────────────────────────────
# Calibrated against the one real (Sonnet-generated) analysis + script pair
# stored in the repo — storage/analysis/t/chapter_0.json and its episode script,
# 23 concept topics, 14 segments, 2,835 words. Measured with the rules above:
#
#   the script as generated                              0.913   (2 honest misses)
#   its last third of segments deleted                   0.870
#   its middle third of segments deleted                 0.870
#   a third of the TOPICS never mentioned                0.652
#   half the script's segments deleted                   0.609
#   two thirds of its segments deleted                   0.348
#
# Note which axis moves. Deleting a third of the SEGMENTS barely moves the
# number (0.913 → 0.870) because a good script returns to its concepts; what
# moves it is topics genuinely never mentioned, which is exactly the failure
# this gate exists to catch, and which lands at ~0.65 by construction.
#
# 0.75 is the shortfall line: 0.163 below the one healthy script measured
# (≈ 3.7 more misses than its 2 out of 23) and 0.10 clear of the 0.65 a
# third-dropped script produces. Both margins are 2-4 topics on a 23-topic
# chapter, which is as much air as one sample entitles anyone to claim.
_COVERAGE_MIN = 0.75

# 0.30 is the floor, and it is a different claim: not "thin" but "not about this
# chapter". The worst degradation reproducible from real output — two thirds of
# the segments deleted — measured 0.348, so nothing a merely-thin script can do
# reaches here. Below 0.30, seven topics in ten are never mentioned; that is the
# same class of defect as the zero-video-parts case worker/process.py already
# raises on, so it raises too. It raises BEFORE slides, TTS and render, so the
# waste is one script call rather than a whole kit.
_COVERAGE_FLOOR = 0.30

# …but only when there are enough topics for the ratio to mean anything. With
# 3 topics a single honest miss reads as 0.67 and two as 0.33 — a good short
# lesson would trip the floor on rounding. At 6 the floor needs 5 of 6 topics
# missing before it fires. Under this count the number is still MEASURED and
# RECORDED (it is a data point for the Sonnet-vs-Haiku comparison either way);
# it is just not allowed to fail a teacher's job.
_MIN_GATED_TOPICS = 6

# How many missed topic names to persist. Enough to diagnose from the row alone,
# short enough that generations.params stays small.
_MAX_RECORDED_MISSES = 12


def gate_mode() -> str:
    """``off`` | ``warn`` (default) | ``strict``.

    Ships as **warn**: production has 118 completed generations and a coverage
    distribution for exactly none of them, and hard-failing on a number nobody
    has measured is how a monitoring feature becomes an outage — a failed job
    costs the teacher a credit. warn still raises at _COVERAGE_FLOOR, which is
    below anything real output has been shown to produce. Once a week of Sonnet
    baselines is sitting in generations.params.coverage, raising the bar to
    _COVERAGE_MIN is one env var (``COVERAGE_GATE=strict``) — which is the whole
    point of measuring first and thresholding second.
    """
    mode = os.getenv("COVERAGE_GATE", "").strip().lower()
    return mode if mode in ("off", "warn", "strict") else "warn"


# ── Arabic-script folding ────────────────────────────────────────────────────
# The Persian incident of 2026-08-16 (issue 8b79d4e0): a Persian book detected
# as "ar" measured 0.059 on its first draft because Arabic-script text has TWO
# ways of writing the same word that byte-compare as different:
#
#   HARAKAT — the optional vowel marks (U+064B–U+065F, plus the superscript
#             alef U+0670). A vocalised topic name "مَدْرَسَة" and the plain
#             "مدرسة" every script writes are the same word; kept as token
#             characters (the Devanagari rule below) they never match.
#             Tatweel (U+0640) is the same problem in Lm clothing — a purely
#             typographic stretch character that isalnum() happily glues in.
#
#   VARIANTS — Persian/Urdu reuse the Arabic block with their own codepoints
#             for the same letters: ک (U+06A9) vs ك (U+0643), ی (U+06CC) vs
#             ي (U+064A). An analyzer that names topics in one convention and
#             a script generator that writes the other match ~nothing, which
#             is exactly what 0.059 looks like.
#
# The fold is Arabic-block ONLY, by codepoint, so the "keep combining marks"
# rule in _tokens below — which Devanagari and Telugu tokenisation depends on —
# is untouched: U+094D and Telugu vowel signs are nowhere in this table.
# Applied to BOTH sides by construction, since topics and haystack alike go
# through _tokens (and _Haystack folds its substring text the same way).
_ARABIC_FOLD = {
    **{cp: None for cp in range(0x064B, 0x0660)},  # harakat: fathatan…wavy hamza
    0x0670: None,   # superscript alef (the mark in قُرْآن-style spellings)
    0x0640: None,   # tatweel — typographic stretch, category Lm so isalnum-true
    0x06A9: 0x0643,  # ک Persian kaf      → ك  (گ gaf is a REAL distinct letter; stays)
    0x06CC: 0x064A,  # ی Persian yeh      → ي
    0x0649: 0x064A,  # ى alef maksura     → ي  (word-final yeh convention)
    0x06D2: 0x064A,  # ے Urdu barree yeh  → ي
    0x06C1: 0x0647,  # ہ Urdu heh goal    → ه
    0x06C2: 0x0647,  # ۂ heh goal + hamza → ه
    0x06C3: 0x0647,  # ۃ teh marbuta goal → ه
    0x0629: 0x0647,  # ة teh marbuta      → ه  (Persian writes it ه: مدرسة/مدرسه)
}


def _fold_arabic(text: str) -> str:
    """Normalise Arabic-script text so spelling conventions compare equal."""
    return text.translate(_ARABIC_FOLD)


def _tokens(text: str) -> list[str]:
    """Casefolded word tokens, in ANY script this product's books are written in.

    Deliberately NOT a ``\\w+`` / ``[^\\W\\d_]+`` regex. Python's ``\\w`` is
    ``str.isalnum()``, and Devanagari and Telugu vowel signs are categories Mc/Mn
    — not alphanumeric — so the regex shatters "प्रकाश संश्लेषण" into 15 one- and
    two-character fragments and "కిరణజన్య సంయోగక్రియ" into 12. A gate built on
    that would report near-zero coverage for every Hindi, Marathi and Telugu
    lesson we ship. Combining marks are therefore kept as part of their token.
    (agent5_slides/figures._words has the same blind spot in a narrower form —
    its ``[a-z][a-z0-9]{2,}`` returns nothing at all outside Latin — which is
    why it is not reused here.)

    Pure digits are dropped (an exercise number is not a topic), and a trailing
    ASCII "s" is stripped from tokens longer than three characters — the same
    crude singularisation figures._words uses, which also covers the -s plurals
    of the Malay/Spanish/Portuguese/French locales and is a no-op in the scripts
    that have no ASCII letters at all.

    Arabic-script text is folded first (_ARABIC_FOLD above): harakat and
    tatweel stripped, Persian/Urdu letter variants mapped to their Arabic
    twins. Every other script passes through the fold untouched.
    """
    out: list[str] = []
    cur: list[str] = []
    for ch in _fold_arabic(text or ""):
        if ch.isalnum() or unicodedata.category(ch)[0] == "M":
            cur.append(ch)
        elif cur:
            out.append("".join(cur))
            cur = []
    if cur:
        out.append("".join(cur))

    toks: list[str] = []
    for raw in out:
        tok = raw.casefold()
        if tok.isdigit():
            continue
        if len(tok) > 3 and tok.endswith("s"):
            tok = tok[:-1]
        toks.append(tok)
    return toks


def _is_spaceless(token: str) -> bool:
    """Whether a token is written in a script that doesn't separate words.

    Chinese, Japanese and Thai text tokenises as one run per punctuation break,
    so set membership can never match a topic name inside it — those tokens are
    matched by SUBSTRING instead. None of the six lesson languages is spaceless,
    but uploaded books are (the 1,008-page Japanese scan of 2026-08-09), and a
    whole-locale false alarm is worth four lines to avoid.
    """
    return any(script_of(ch) in SPACELESS_SCRIPTS for ch in token)


class _Haystack:
    """The artifact's text, indexed once for repeated topic lookups."""

    __slots__ = ("text", "tokens", "stems")

    def __init__(self, text: str) -> None:
        toks = _tokens(text)
        # Folded like the tokens are, so the spaceless-script substring path in
        # has() compares like with like — a topic key has already been through
        # _fold_arabic by the time it arrives here.
        self.text = _fold_arabic(text or "").casefold()
        self.tokens = set(toks)
        self.stems = {t[:_STEM] for t in toks}

    def has(self, token: str) -> bool:
        if token in self.tokens:
            return True
        if len(token) >= _STEM and token[:_STEM] in self.stems:
            return True
        # Spaceless scripts only — see _is_spaceless. Substring, because the
        # whole sentence is one token.
        return _is_spaceless(token) and token in self.text


def _pooled_denominator(analysis: dict, episode: dict | None) -> bool:
    """Whether topic scoping falls back to ALL of the analysis's concepts.

    True when an episode declares no concept ids and the plan has at most one
    episode — the empty-ids fallback in chapter_topics below. The claim that
    fallback rests on ("a single-episode plan covers the whole chapter") is
    TRUE for whole-chapter jobs and FALSE for chapter-PART jobs, which is why
    measure() needs to know the fallback fired: a part job's analysis also
    yields a single-episode plan with empty ids, and pooling there judges a
    part-3 script against topics parts 1, 2 and 4 own (live incident 8b79d4e0
    — a part 3/4 hard-failed at "3 of 17 topics").
    """
    if episode is None:
        return False
    if episode.get("key_concepts_introduced"):
        return False
    plan = (analysis.get("episodes") or {}).get("episodes") or []
    return len(plan) <= 1


def chapter_topics(analysis: dict, episode: dict | None = None) -> list[dict]:
    """The topics an artifact is expected to address, as ``[{origin, name}]``.

    ``episode`` scopes to one part (the presentation path renders one script per
    part). ``None`` means the whole chapter — what a teacher document is
    generated from, since documents are per CHAPTER and never per part.
    """
    topics: list[dict] = []
    seen: set[str] = set()
    chapter_title = str(analysis.get("chapter_title") or "")

    concepts = (analysis.get("concepts") or {}).get("concepts") or []
    wanted = set(episode.get("key_concepts_introduced") or []) if episode is not None else None
    if wanted == set() and _pooled_denominator(analysis, episode):
        # A single-episode plan that declares no concept ids covers the whole
        # chapter by definition — the pre-chunking analyzer and segmenter.py
        # both leave key_concepts_introduced empty ("populated later"), and a
        # cached v1 row still looks like that. Scoping to an empty set would
        # turn the gate off for exactly those rows instead of measuring them.
        # (For a chapter-PART job the "by definition" is wrong — measure()
        # detects that via _pooled_denominator and marks the report pooled.)
        wanted = None
    for concept in concepts:
        if not isinstance(concept, dict):
            continue
        if wanted is not None and concept.get("concept_id") not in wanted:
            continue
        name = " ".join(str(concept.get("name") or "").split())
        if name and name.casefold() not in seen:
            seen.add(name.casefold())
            topics.append({"origin": "concept", "name": name})

    if episode is not None:
        section_lists = [episode.get("sections_covered") or []]
    else:
        section_lists = [
            ep.get("sections_covered") or []
            for ep in ((analysis.get("episodes") or {}).get("episodes") or [])
        ]
    for titles in section_lists:
        for raw in titles:
            name = " ".join(str(raw).split())
            key = name.casefold()
            # A heading that names nothing is a heading nothing can be measured
            # against — one rule, shared with the display path.
            if key in seen or names_nothing(name, chapter_title):
                continue
            seen.add(key)
            topics.append({"origin": "section", "name": name})

    return topics


def _topic_keys(topics: list[dict]) -> list[list[str]]:
    """The distinguishing tokens of each topic, rarest-first and capped."""
    raw = [_tokens(t["name"]) for t in topics]
    df: dict[str, int] = {}
    for toks in raw:
        for tok in set(toks):
            df[tok] = df.get(tok, 0) + 1

    keys: list[list[str]] = []
    for toks in raw:
        distinct = [t for t in toks if df[t] < _DROP_TOKEN_AT_DF] or toks
        # Rarest first so the cap keeps the discriminating words; the length
        # tie-break prefers the longer (more specific) of two equally rare ones.
        distinct = sorted(dict.fromkeys(distinct), key=lambda t: (df[t], -len(t)))
        keys.append(distinct[:_MAX_TOPIC_KEYS])
    return keys


def measure(analysis: dict, episode: dict | None, produced_text: str, *,
            part_scoped: bool = False) -> dict:
    """Measure topic coverage of ``produced_text``. Pure; no model call.

    Returns ``{checked, covered, topics, addressed, missed, sections, concepts,
    gated, verdict}``. ``checked`` is False when there is nothing measurable to
    compare against (an analysis that extracted no concepts and whose only
    heading is a placeholder), and then ``covered`` is None and means nothing —
    the same principle book_health.text_quality applies to a book with too
    little text to judge. Unjudgeable is not the same as bad.

    ``part_scoped`` says the artifact was generated from ONE PART of a chapter
    (worker/process.py's params.part path). When such a job's denominator had
    to be pooled to the whole concept list (_pooled_denominator), the topics
    include what OTHER parts teach, so the ratio measures the chapter split,
    not the script. The report then carries ``"pooled": true`` and a verdict
    that cannot fail the job — the number is still recorded (it is a data
    point for the model comparison either way), it just stops pretending to be
    a coverage judgment. Whole-chapter callers never set this, and pooling for
    them keeps failing at the floor exactly as before.
    """
    pooled = part_scoped and _pooled_denominator(analysis, episode)
    topics = chapter_topics(analysis, episode)
    keys = _topic_keys(topics)
    hay = _Haystack(produced_text)

    addressed: list[dict] = []
    missed: list[dict] = []
    for topic, topic_keys in zip(topics, keys):
        if not topic_keys:
            continue  # unmeasurable (a name with no word characters at all)
        hits = sum(1 for key in topic_keys if hay.has(key))
        (addressed if hits / len(topic_keys) >= _ADDRESSED_SHARE else missed).append(topic)

    total = len(addressed) + len(missed)
    out = {
        "checked": total > 0,
        "covered": None,
        "topics": total,
        "addressed": len(addressed),
        "sections": sum(1 for t in addressed + missed if t["origin"] == "section"),
        "concepts": sum(1 for t in addressed + missed if t["origin"] == "concept"),
        "missed": [t["name"] for t in missed][:_MAX_RECORDED_MISSES],
        "gated": total >= _MIN_GATED_TOPICS,
        "verdict": "unmeasured",
    }
    if pooled:
        out["pooled"] = True
    if not total:
        return out

    covered = len(addressed) / total
    out["covered"] = round(covered, 3)
    if covered < _COVERAGE_FLOOR:
        out["verdict"] = "floor"
    elif covered < _COVERAGE_MIN:
        out["verdict"] = "short"
    else:
        out["verdict"] = "ok"
    if pooled and out["verdict"] in ("short", "floor"):
        # Not a failing verdict, because it is not a failure claim: the missed
        # list is largely other parts' topics. "pooled" in the row is what lets
        # the analytics query exclude these ratios from the model comparison —
        # or study them, since they measure how a chapter's concepts split
        # across its parts.
        out["verdict"] = "pooled"
    return out


def should_fail(report: dict, mode: str | None = None) -> bool:
    """Whether this report may fail the job, under the deployment's gate mode.

    A report that isn't gated (too few measurable topics — see
    _MIN_GATED_TOPICS) never fails one, whatever the mode. Neither does a
    pooled part-scoped report (see measure): its denominator is the whole
    chapter's concept list, so a part-3 script "failing" it is the gate
    failing, not the script — that is the belt to the verdict rewrite's
    braces, so a pooled report stays harmless even if a future verdict
    change forgets about it.
    """
    mode = mode or gate_mode()
    if mode == "off" or not report.get("gated") or not report.get("checked"):
        return False
    if report.get("pooled"):
        return False
    if mode == "strict":
        return report["verdict"] in ("short", "floor")
    return report["verdict"] == "floor"


def should_retry(report: dict, mode: str | None = None) -> bool:
    """Whether a first draft has earned the one must_cover retry.

    Same bar as should_fail — the retry exists to rescue a job that would
    otherwise fail outright, so it never fires below that bar — MINUS pooled
    part-scoped reports, and that exclusion is the point of this function
    existing separately: a pooled report's ``missed`` list is mostly OTHER
    parts' topics, so a retry built from it explicitly orders the model to
    teach material this part does not contain. That is the actively harmful
    half of incident 8b79d4e0 — the part-3 must_cover retry told the script
    to cover parts 1, 2 and 4. Kept here rather than inline in the worker so
    the decision is testable without importing worker.process.
    """
    return should_fail(report, mode) and not report.get("pooled")


def script_text(script: dict) -> str:
    """Every word an episode script puts in front of a student — spoken AND
    on-screen.

    Slide text counts: a concept introduced only as a slide heading or a diagram
    node has still been taught, and scoring the narration alone would mark those
    as missed. ``elevenlabs_text`` is deliberately skipped — it is ``text`` plus
    SSML break tags, so including it would double every token for nothing.
    """
    out: list[str] = []
    for seg in script.get("segments") or []:
        if not isinstance(seg, dict):
            continue
        out.append(seg.get("text") or "")
        out.append(seg.get("slide_heading") or "")
        out.extend(str(p) for p in (seg.get("slide_points") or []))
        visual = seg.get("slide_visual") or {}
        if isinstance(visual, dict):
            out.extend(str(n) for n in (visual.get("nodes") or []))
            out.extend(str(o) for o in (visual.get("options") or []))
            out.append(str(visual.get("caption") or ""))
            out.append(str(visual.get("body") or ""))
            for group in visual.get("groups") or []:
                if isinstance(group, dict):
                    out.append(str(group.get("heading") or ""))
                    out.extend(str(i) for i in (group.get("items") or []))
            for item in visual.get("items") or []:
                if isinstance(item, dict):
                    out.append(str(item.get("label") or ""))
    return " ".join(p for p in out if p)


def docx_text(path) -> str:
    """All text in a built .docx — paragraphs AND table cells.

    Tables are not optional: a worksheet's matching exercise and an exam's mark
    scheme live entirely inside them, so a paragraph-only read would score those
    kinds as empty. Best-effort by contract — the caller treats a failure here as
    "not measured", never as a coverage failure.
    """
    import docx  # local: only the document kinds need it

    document = docx.Document(str(path))
    out = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            out.extend(cell.text for cell in row.cells)
    return " ".join(t for t in out if t)
