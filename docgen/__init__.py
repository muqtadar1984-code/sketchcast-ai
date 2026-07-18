"""Document generators — turn a chapter into an editable .docx (lesson plan,
class activities, exam paper). Dispatched from the worker by generation `kind`."""

from __future__ import annotations

from pathlib import Path


class _DirectiveClient:
    """Client proxy that appends a language directive to every prompt — one
    place covers all five document builders (and future ones) without touching
    their individual prompt construction."""

    def __init__(self, inner, directive: str):
        self._inner = inner
        self._directive = directive

    def analyze(self, prompt, **kwargs):
        return self._inner.analyze(prompt + self._directive, **kwargs)

    def __getattr__(self, name):
        return getattr(self._inner, name)


def generate_document(
    kind: str,
    book: dict,
    chapter: dict,
    analysis: dict,
    client,
    params: dict,
    out_dir: Path,
    template: str | None = None,
    language: str = "en",
) -> Path:
    """Build the .docx for `kind` and return its path. `template` = optional
    school .docx whose styles/header/footer/logo the document inherits."""
    from shared.languages import is_rtl, prompt_directive

    directive = prompt_directive(language)
    if directive:
        client = _DirectiveClient(client, directive)
    builders = {
        "lesson_plan": "docgen.lesson_plan",
        "activity": "docgen.activity",
        "exam_paper": "docgen.exam_paper",
        "worksheet": "docgen.worksheet",
        "case_study": "docgen.case_study",
    }
    mod_name = builders.get(kind)
    if not mod_name:
        raise ValueError(f"Unknown document kind: {kind}")
    import importlib
    build = importlib.import_module(mod_name).build
    # Every generator takes the same signature.
    path = build(book, chapter, analysis, client, params or {}, out_dir, template)
    # RTL (Arabic) documents: Word does the SHAPING itself, but paragraphs
    # need right alignment + the bidi flag. A post-pass over the saved file
    # keeps all five builders untouched. Best-effort — never fails the doc.
    if is_rtl(language):
        try:
            _mirror_docx_rtl(path)
        except Exception:  # noqa: BLE001
            import logging

            logging.getLogger("docgen").warning("RTL docx pass failed for %s", path)
    return path


def _mirror_docx_rtl(docx_path) -> None:
    import docx as _docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    d = _docx.Document(str(docx_path))

    def _mirror_paragraph(p) -> None:
        # Centred paragraphs (titles, badges) keep their centring.
        if p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pPr = p._p.get_or_add_pPr()  # noqa: SLF001
        if pPr.find(qn("w:bidi")) is None:
            bidi = pPr.makeelement(qn("w:bidi"), {})
            pPr.append(bidi)

    for p in d.paragraphs:
        _mirror_paragraph(p)
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _mirror_paragraph(p)
    d.save(str(docx_path))
