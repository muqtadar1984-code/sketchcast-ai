"""python-docx helpers + the per-kind STYLE LAYER shared by the document generators.

Three founder-approved visual styles (mockups: scratchpad docx-styles, 2026-08):

  · "exam_board"  — formal serif B/W (Georgia/Times), rule-bound masthead,
                    examiner marks table, small-caps section heads, right-tabbed
                    [n] marks, ruled T/F answer boxes, dotted writing lines.
                    Used by exam_paper.py and exam.py (paper + answer key).
  · "workbook"    — colorful Calibri, teal/amber/violet section chips, callout
                    instructions, tinted answer panels, teal page-number square.
                    Used by worksheet.py and activity.py.
  · "classroom"   — friendly-professional, single teal accent, full-width
                    section bars. Used by lesson_plan.py and case_study.py.

The builders keep ALL their content logic (question assembly, renderer-owned
numbering, answer keys); they only call the semantic helpers below, which
dispatch on the document's style.

Direction-awareness: every helper honors the lesson language's direction —
RTL (Arabic / Jawi) documents get w:bidi paragraphs, mirrored hanging indents,
right-aligned chips/tables (w:bidiVisual), swapped footer cells, complex-script
fonts (w:rFonts w:cs + w:szCs — Georgia/Times have no Arabic glyphs), and NO
letter-spacing (expanded spacing breaks Arabic glyph joining).

School letterhead (new_doc(template=...)): the branded template owns the page
masthead and footer, so the style DEMOTES its own — no title rules/two-tone
bars, no SketchCast footer — while section styling (chips/bars/tables) stays.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# User-visible strings + enumeration letters, localized for every lesson
# language, live in docgen/strings.py (one table, one _t). Re-exported here so
# the builders keep their single `docgen import docx_builder as dx` import.
from docgen.strings import (  # noqa: F401  (re-exports used by builders)
    RTL_LANGS as _RTL_LANGS,
    _STRINGS,
    _t,
    column_label,
    difficulty_label,
    letters,
    marks_bracket,
    match_instruction,
    section_heading,
    tf_word,
)

# Legacy palette (kept for compatibility with older imports).
GREEN = RGBColor(0x2E, 0x6B, 0x4E)
GREY = RGBColor(0x6F, 0x6A, 0x5F)


# ── Style registry ───────────────────────────────────────────────────────────

@dataclass(frozen=True)
class _Style:
    name: str
    body_font: str
    head_font: str
    body_size: float          # pt
    ink: str                  # heading text hex
    body_color: str           # body text hex
    meta_color: str           # subtitle / footer grey
    line_color: str           # writing-line / hairline grey
    accents: tuple = ()       # (fill, on_fill_text, dark_text, tint) cycles


# workbook palette (mockup: build_modern_workbook.py)
_WB_TEAL = ("17A398", "FFFFFF", "0E7A70", "E2F4F1")
_WB_AMBER = ("F5A623", "4A3800", "9A6500", "FDF2DE")
_WB_VIOLET = ("7C6FD0", "FFFFFF", "5647B5", "EDEAF9")

# classroom palette (mockup: build_classroom_clean.py)
_CC_TEAL = ("1FB8A6", "FFFFFF", "0E7A70", "E2F4F1")
_CC_SLATE = ("14181F", "FFFFFF", "14181F", "F0F3F5")

STYLES: dict[str, _Style] = {
    "exam_board": _Style(
        name="exam_board", body_font="Georgia", head_font="Times New Roman",
        body_size=11, ink="000000", body_color="000000", meta_color="000000",
        line_color="000000", accents=(("000000", "FFFFFF", "000000", "F2F2F2"),),
    ),
    "workbook": _Style(
        name="workbook", body_font="Calibri", head_font="Calibri",
        body_size=11, ink="263238", body_color="333333", meta_color="6B7778",
        line_color="AEBDBB", accents=(_WB_TEAL, _WB_AMBER, _WB_VIOLET),
    ),
    "classroom": _Style(
        name="classroom", body_font="Calibri", head_font="Calibri Light",
        body_size=11, ink="14181F", body_color="212730", meta_color="6B727C",
        line_color="C6CCD4", accents=(_CC_TEAL,),
    ),
}

STYLE_BY_KIND: dict[str, str] = {
    "exam_paper": "exam_board",
    "exam": "exam_board",
    "worksheet": "workbook",
    "activity": "workbook",
    "lesson_plan": "classroom",
    "case_study": "classroom",
}

_HAIR = "D9DEE4"       # hairline table border (classroom)
_DOT_GREY = "C7CFCE"   # dotted separators (workbook)

# Complex-script (w:cs) font per script family: Georgia/Times/Calibri carry no
# Arabic — Word needs an explicit cs font or it substitutes unpredictably.
_CS_FONTS = {
    "exam_board": {"ar": "Traditional Arabic", "ms-arab": "Traditional Arabic",
                   "hi": "Nirmala UI", "mr": "Nirmala UI", "te": "Nirmala UI"},
    "_default": {"ar": "Arial", "ms-arab": "Arial",
                 "hi": "Nirmala UI", "mr": "Nirmala UI", "te": "Nirmala UI"},
}


# ── Per-document context ─────────────────────────────────────────────────────

@dataclass
class _Ctx:
    style: _Style
    lang: str = "en"
    rtl: bool = False
    branded: bool = False       # school letterhead template present
    accent_i: int = -1          # cycles per level-1 section (workbook)
    width: int = 9412           # content width in dxa (recomputed per doc)

    @property
    def accent(self):
        accents = self.style.accents
        return accents[max(self.accent_i, 0) % len(accents)]


_CTX_ATTR = "_sketchcast_style_ctx"


def _ctx(doc: Document) -> _Ctx:
    ctx = getattr(doc, _CTX_ATTR, None)
    if ctx is None:  # helper used on a doc that didn't come from new_doc()
        ctx = _Ctx(style=STYLES["classroom"])
        setattr(doc, _CTX_ATTR, ctx)
    return ctx


# ── Low-level XML helpers ────────────────────────────────────────────────────

# Schema-ordered successors of w:pBdr inside w:pPr (python-docx has no API).
_PBDR_SUCCESSORS = (
    "w:shd", "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap",
    "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN",
    "w:bidi", "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
    "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
    "w:textDirection", "w:textAlignment", "w:textboxTightWrap",
    "w:outlineLvl", "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange",
)


def _p_border(p, edge: str, val: str = "single", sz: int = 4, space: int = 1,
              color: str = "000000") -> None:
    pPr = p._p.get_or_add_pPr()  # noqa: SLF001
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.insert_element_before(pBdr, *_PBDR_SUCCESSORS)
    el = pBdr.find(qn("w:" + edge))
    if el is None:
        el = OxmlElement("w:" + edge)
        pBdr.append(el)
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)


def _p_bidi(p) -> None:
    pPr = p._p.get_or_add_pPr()  # noqa: SLF001
    if pPr.find(qn("w:bidi")) is None:
        bidi = OxmlElement("w:bidi")
        pPr.insert_element_before(
            bidi, "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
            "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap",
            "w:jc", "w:textDirection", "w:textAlignment", "w:textboxTightWrap",
            "w:outlineLvl", "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr",
            "w:pPrChange")


def _letter_space(run, pts: float) -> None:
    """Expanded character spacing. NEVER call for RTL text — expanded spacing
    breaks Arabic glyph joining (letters stop connecting)."""
    rPr = run._r.get_or_add_rPr()  # noqa: SLF001
    el = OxmlElement("w:spacing")
    el.set(qn("w:val"), str(int(pts * 20)))
    rPr.insert_element_before(
        el, "w:w", "w:kern", "w:position", "w:sz", "w:szCs", "w:highlight",
        "w:u", "w:effect", "w:bdr", "w:shd", "w:fitText", "w:vertAlign",
        "w:rtl", "w:cs", "w:em", "w:lang", "w:eastAsianLayout",
        "w:specVanish", "w:oMath")


def _caps(run) -> None:
    """Display-only ALL CAPS (w:caps): the run's text stays untouched, so the
    tests' heading strings keep their exact current text."""
    rPr = run._r.get_or_add_rPr()  # noqa: SLF001
    caps = OxmlElement("w:caps")
    rPr.insert_element_before(
        caps, "w:smallCaps", "w:strike", "w:dstrike", "w:outline", "w:shadow",
        "w:emboss", "w:imprint", "w:noProof", "w:snapToGrid", "w:vanish",
        "w:webHidden", "w:color", "w:spacing", "w:w", "w:kern", "w:position",
        "w:sz", "w:szCs", "w:highlight", "w:u", "w:effect", "w:bdr", "w:shd",
        "w:fitText", "w:vertAlign", "w:rtl", "w:cs", "w:em", "w:lang",
        "w:eastAsianLayout", "w:specVanish", "w:oMath")


def _cs_font_for(ctx: _Ctx) -> str:
    table = _CS_FONTS.get(ctx.style.name) or _CS_FONTS["_default"]
    fallback = _CS_FONTS["_default"]
    return table.get(ctx.lang) or fallback.get(ctx.lang) or ctx.style.body_font


def _run(p, text: str, ctx: _Ctx, *, size: float | None = None,
         font: str | None = None, bold: bool = False, italic: bool = False,
         color: str | None = None, small_caps: bool = False,
         spacing: float | None = None, caps: bool = False,
         rtl: bool | None = None):
    """Add a run with the style's fonts + complex-script (cs) handling."""
    run = p.add_run(text)
    f = font or ctx.style.body_font
    sz = size if size is not None else ctx.style.body_size
    run.font.name = f
    run.font.size = Pt(sz)
    run.font.bold = bold
    run.font.italic = italic
    run.font.small_caps = small_caps
    run.font.color.rgb = RGBColor.from_string(color or ctx.style.body_color)
    rPr = run._r.get_or_add_rPr()  # noqa: SLF001
    rfonts = rPr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), f)
    rfonts.set(qn("w:hAnsi"), f)
    # cs font + size: Arabic (and Indic) glyphs come from the cs font — the
    # exam style's Georgia/Times has none, so Word must be told explicitly.
    cs = _cs_font_for(ctx)
    rfonts.set(qn("w:cs"), cs)
    szCs = OxmlElement("w:szCs")
    # Traditional Arabic renders visually small at equal pt — bump it slightly.
    cs_pt = sz + 2 if cs == "Traditional Arabic" else sz
    szCs.set(qn("w:val"), str(int(cs_pt * 2)))
    old = rPr.find(qn("w:szCs"))
    if old is not None:
        rPr.remove(old)
    rPr.insert_element_before(
        szCs, "w:highlight", "w:u", "w:effect", "w:bdr", "w:shd", "w:fitText",
        "w:vertAlign", "w:rtl", "w:cs", "w:em", "w:lang", "w:eastAsianLayout",
        "w:specVanish", "w:oMath")
    if caps:
        _caps(run)
    if spacing and not ctx.rtl:  # no letter-spacing on RTL scripts
        _letter_space(run, spacing)
    # RUN-LEVEL direction: w:bidi on the paragraph mirrors the BLOCK, but the
    # line's base ordering stays LTR without <w:rtl/> on each run — question
    # numbers ("1.") would render at the visual LEFT (read last) and digit
    # brackets reverse ("]1[") in Arabic. Verified by rendering: adding w:rtl
    # flips the same paragraph to native RTL (number right, period left).
    if ctx.rtl if rtl is None else rtl:
        el = rPr.find(qn("w:rtl"))
        if el is None:
            el = OxmlElement("w:rtl")
            rPr.insert_element_before(el, "w:cs", "w:em", "w:lang",
                                      "w:eastAsianLayout", "w:specVanish", "w:oMath")
    return run


def _ensure_style(doc: Document, name: str):
    """doc.styles[name], fabricating the style when the base document lacks it.

    Branded documents build on a teacher's UPLOADED letterhead .docx, and Word
    keeps its built-in styles LATENT until a document actually uses them — so
    a perfectly normal letterhead defines no 'Title'/'Heading N', python-docx
    cannot see latent styles, and doc.styles['Title'] raises KeyError
    (2026-08-18 prod failure: every document for a letterhead user died with
    "no style with name 'Title'"). Fabricate a minimal definition instead:
    same name, based on Normal (mandatory in every valid docx), outline level
    for Heading N so Word's navigation pane still sees the headings. Every
    visual property is direct-formatted by the callers, so the fabricated
    style needs no look of its own.
    """
    try:
        return doc.styles[name]
    except KeyError:
        pass
    st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH, builtin=True)
    try:
        st.base_style = doc.styles["Normal"]
    except KeyError:
        pass  # a styles part without Normal is broken beyond our care
    st.quick_style = True
    m = re.fullmatch(r"Heading (\d)", name)
    if m:
        # w:outlineLvl is what the navigation pane reads; python-docx has no API.
        ppr = st.element.get_or_add_pPr()
        lvl = OxmlElement("w:outlineLvl")
        lvl.set(qn("w:val"), str(int(m.group(1)) - 1))
        ppr.append(lvl)
    return st


def _p(doc: Document, ctx: _Ctx, *, align=None, before: float = 0,
       after: float = 0, left=None, hang=None, line=None, style: str | None = None):
    """New body paragraph with direction handling.

    RTL: only w:bidi is set — NO explicit w:jc. Word interprets jc left/right
    LOGICALLY under bidi (they swap: jc=right renders LEFT — verified by
    rendering), while a bare bidi paragraph defaults to start = physical
    right, which is exactly the mirrored layout. The same applies to
    w:ind left/hanging, so hanging indents mirror to the right edge for free.
    Callers that need a specific physical side pass `align` themselves and
    account for the swap."""
    par = doc.add_paragraph(style=_ensure_style(doc, style) if style else None)
    pf = par.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if left is not None:
        pf.left_indent = left
    if hang is not None:
        pf.first_line_indent = -hang
    if line is not None:
        pf.line_spacing = line
    if align is not None:
        par.alignment = align
    if ctx.rtl:
        _p_bidi(par)
    return par


def _fmt_cell_par(cell, ctx: _Ctx, align=None):
    """Cell paragraph. RTL gets w:bidi and NO jc (see _p: bidi's default start
    alignment is the physical right; explicit jc values swap under bidi)."""
    par = cell.paragraphs[0]
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(0)
    if align is not None:
        par.alignment = align
    if ctx.rtl:
        _p_bidi(par)
    return par


def _cell_extra_par(cell, ctx: _Ctx):
    par = cell.add_paragraph()
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(0)
    if ctx.rtl:
        _p_bidi(par)
    return par


def _shade_cell(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()  # noqa: SLF001
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _cell_margins(cell, top=None, left=None, bottom=None, right=None) -> None:
    tcPr = cell._tc.get_or_add_tcPr()  # noqa: SLF001
    mar = OxmlElement("w:tcMar")
    for name, v in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        if v is not None:
            el = OxmlElement("w:" + name)
            el.set(qn("w:w"), str(v))
            el.set(qn("w:type"), "dxa")
            mar.append(el)
    tcPr.append(mar)


def _cell_borders(cell, **edges) -> None:
    """edges: top/left/bottom/right -> None (=nil) or (val, sz, color)."""
    tcPr = cell._tc.get_or_add_tcPr()  # noqa: SLF001
    tcB = tcPr.find(qn("w:tcBorders"))
    if tcB is None:
        tcB = OxmlElement("w:tcBorders")
        tcPr.append(tcB)
    for edge in ("top", "left", "bottom", "right"):
        if edge not in edges:
            continue
        spec = edges[edge]
        el = tcB.find(qn("w:" + edge))
        if el is None:
            el = OxmlElement("w:" + edge)
            tcB.append(el)
        if spec is None:
            el.set(qn("w:val"), "nil")
        else:
            val, sz, color = spec
            el.set(qn("w:val"), val)
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)


def _table_borders(tbl, sz: int = 4, color: str = "000000", none: bool = False) -> None:
    tblPr = tbl._tbl.tblPr  # noqa: SLF001
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        if none:
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
        else:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color if not none else "auto")
        borders.append(el)
    tblPr.append(borders)


def _table_fixed(tbl, widths_dxa: list[int]) -> None:
    tbl.autofit = False
    tblPr = tbl._tbl.tblPr  # noqa: SLF001
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths_dxa)))
    tblW.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    for row in tbl.rows:
        for i, w in enumerate(widths_dxa[: len(row.cells)]):
            row.cells[i].width = w * 635  # dxa → EMU
    for i, w in enumerate(widths_dxa[: len(tbl.columns)]):
        tbl.columns[i].width = w * 635


def _table_bidi(tbl, ctx: _Ctx) -> None:
    """RTL tables: w:bidiVisual flips the column order so 'Column A' reads
    from the right, matching the mirrored page."""
    if not ctx.rtl:
        return
    tblPr = tbl._tbl.tblPr  # noqa: SLF001
    if tblPr.find(qn("w:bidiVisual")) is None:
        bidi = OxmlElement("w:bidiVisual")
        tblPr.insert(0, bidi)
        # bidiVisual must precede tblW/jc/etc.; tblStyle (if any) must stay first
        style_el = tblPr.find(qn("w:tblStyle"))
        if style_el is not None:
            tblPr.remove(bidi)
            style_el.addnext(bidi)


def _row_height(row, twips: int, rule: str = "atLeast") -> None:
    trPr = row._tr.get_or_add_trPr()  # noqa: SLF001
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(twips))
    h.set(qn("w:hRule"), rule)
    trPr.append(h)


def _page_field(p, ctx: _Ctx, instr: str, size: float = 8.5,
                color: str | None = None, bold: bool = False) -> None:
    """A PAGE / NUMPAGES field. Field runs are forced LTR: the counter is a
    digits-and-slash fragment, and RTL runs bidi-scramble "1 / 3" to "13 /"."""
    for part, is_field in (("begin", True), (instr, False), ("end", True)):
        run = _run(p, "", ctx, size=size, color=color or ctx.style.meta_color,
                   bold=bold, rtl=False)
        if is_field:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), part)
            run._r.append(fld)  # noqa: SLF001
        else:
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = f" {part} "
            run._r.append(it)  # noqa: SLF001


def _content_width_dxa(doc: Document) -> int:
    sec = doc.sections[0]
    try:
        return int((sec.page_width - sec.left_margin - sec.right_margin) / 635)
    except TypeError:  # template with missing geometry
        return 9412


# ── Chips / bars / rules ─────────────────────────────────────────────────────

def _bar(doc: Document, ctx: _Ctx, text: str, fill: str, text_color: str,
         *, full_width: bool, heading_style: str, size: float = 12):
    """A shaded single-cell-table section bar/chip (paragraph shading can't
    carry padding or partial width, so the mockups use one-cell tables)."""
    tbl = doc.add_table(rows=1, cols=1)
    _table_borders(tbl, none=True)
    # Table jc also swaps under w:bidiVisual (LEFT renders right, verified by
    # rendering) — so LEFT gives the leading edge in both directions.
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _table_bidi(tbl, ctx)
    cell = tbl.rows[0].cells[0]
    if full_width:
        _table_fixed(tbl, [ctx.width])
    else:
        # auto layout, auto widths: the chip hugs its text whatever the
        # language (add_table pre-sets a full-page grid — undo it).
        tbl.autofit = True
        tblPr = tbl._tbl.tblPr  # noqa: SLF001
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:type"), "auto")
        tblW.set(qn("w:w"), "0")
        for gc in tbl._tbl.tblGrid.findall(qn("w:gridCol")):  # noqa: SLF001
            gc.set(qn("w:w"), "0")
        tcPr = cell._tc.get_or_add_tcPr()  # noqa: SLF001
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is None:
            tcW = OxmlElement("w:tcW")
            tcPr.append(tcW)
        tcW.set(qn("w:type"), "auto")
        tcW.set(qn("w:w"), "0")
    _shade_cell(cell, fill)
    _cell_margins(cell, top=72, left=180, bottom=72, right=180)
    p = _fmt_cell_par(cell, ctx)
    # Real Heading/Title style so Word's navigation (and the tests) still see
    # a heading; every visual property is overridden by direct formatting.
    p.style = _ensure_style(doc, heading_style)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _p_border(p, "bottom", val="nil")  # Title style's built-in rule off
    if ctx.rtl:
        _p_bidi(p)
    _run(p, text, ctx, size=size, bold=True, color=text_color)
    return tbl


def _two_tone_rule(doc: Document, ctx: _Ctx, main_w: int, second_w: int,
                   height: int = 80) -> None:
    """Workbook title underline: teal + amber rules side by side."""
    teal, amber = STYLES["workbook"].accents[0][0], STYLES["workbook"].accents[1][0]
    gap = 130
    tbl = doc.add_table(rows=1, cols=3)
    _table_borders(tbl, none=True)
    _table_fixed(tbl, [main_w, gap, second_w])
    tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT if ctx.rtl else WD_TABLE_ALIGNMENT.LEFT
    _row_height(tbl.rows[0], height, rule="exact")
    fills = (teal, None, amber)
    for cell, fill in zip(tbl.rows[0].cells, fills):
        _cell_margins(cell, top=0, left=0, bottom=0, right=0)
        if fill:
            _shade_cell(cell, fill)
        p = _fmt_cell_par(cell, ctx, align=WD_ALIGN_PARAGRAPH.LEFT)
        r = p.add_run("")
        r.font.size = Pt(1)


def _spacer(doc: Document, pts: float = 4):
    """Tiny empty paragraph: breathing room around tables, and the standard
    fix for Word's border-merging (see writing_lines)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)
    r = p.add_run("")
    r.font.size = Pt(2)
    return p


# ── Document setup ───────────────────────────────────────────────────────────

def _clear_body(doc: Document) -> None:
    """Drop a template's existing paragraphs/tables but keep section properties
    (header/footer references), so generated content sits on the school styling."""
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("}p") or child.tag.endswith("}tbl"):
            body.remove(child)


def _setup_page(doc: Document, style: _Style) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    if style.name == "exam_board":
        sec.top_margin, sec.bottom_margin = Cm(1.9), Cm(2.1)
        sec.left_margin = sec.right_margin = Cm(2.2)
    else:
        sec.top_margin, sec.bottom_margin = Cm(1.9), Cm(2.0)
        sec.left_margin = sec.right_margin = Cm(2.0)
    sec.footer_distance = Cm(1.1)
    normal = doc.styles["Normal"]
    normal.font.name = style.body_font
    normal.font.size = Pt(style.body_size)
    normal.font.color.rgb = RGBColor.from_string(style.body_color)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    rfonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), style.body_font)
    rfonts.set(qn("w:hAnsi"), style.body_font)


def _footer(doc: Document, ctx: _Ctx, brand_text: str) -> None:
    """Brand left, page number right.

    Word gotcha (mockups): custom tab stops added inside the Footer style
    render zero-width — the built-in style's own center/right tabs win — so
    alignment is done with a 2-cell borderless table instead. For RTL the
    table gets w:bidiVisual, which swaps the cells (brand right, page left).
    """
    st = ctx.style
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    tbl = footer.add_table(rows=1, cols=2, width=ctx.width * 635)
    fp._p.addprevious(tbl._tbl)  # noqa: SLF001 — table first, stub paragraph after
    _table_borders(tbl, none=True)
    page_w = 700 if st.name == "workbook" else 1700
    _table_fixed(tbl, [ctx.width - page_w, page_w])
    _table_bidi(tbl, ctx)
    left, right = tbl.rows[0].cells
    if st.name != "workbook":
        # hairline rule above the whole footer line
        for cell in (left, right):
            _cell_borders(cell, top=("single", 4, st.line_color if st.name == "classroom" else "000000"))
        _cell_margins(left, top=40, left=0, bottom=0, right=0)
        _cell_margins(right, top=40, left=0, bottom=0, right=0)
    else:
        _cell_margins(left, top=0, left=0, bottom=0, right=120)
        _cell_margins(right, top=0, left=0, bottom=0, right=0)
    lp = _fmt_cell_par(left, ctx)
    _run(lp, brand_text, ctx, size=8.5, color="7C8888" if st.name == "workbook" else st.meta_color)
    # Page-number cell: numeric LTR content — NO w:bidi (it would reorder
    # "1 / 2"), so the physical alignment below is literal.
    rp = right.paragraphs[0]
    rp.paragraph_format.space_before = Pt(0)
    rp.paragraph_format.space_after = Pt(0)
    rp.alignment = (WD_ALIGN_PARAGRAPH.CENTER if st.name == "workbook"
                    else (WD_ALIGN_PARAGRAPH.LEFT if ctx.rtl else WD_ALIGN_PARAGRAPH.RIGHT))
    if st.name == "workbook":
        # teal page-number square
        acc = st.accents[0]
        _shade_cell(right, acc[0])
        _row_height(tbl.rows[0], 380, rule="atLeast")
        right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _page_field(rp, ctx, "PAGE", size=9, color="FFFFFF", bold=True)
    else:
        # "1 / 2" — numeric, so no per-language wording is needed
        _page_field(rp, ctx, "PAGE")
        _run(rp, " / ", ctx, size=8.5, color=st.meta_color, rtl=False)
        _page_field(rp, ctx, "NUMPAGES")
    # shrink the mandatory trailing footer paragraph to nothing
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    r = fp.add_run("")
    r.font.size = Pt(2)


def new_doc(title: str, subtitle: str = "", template: str | None = None,
            kind: str | None = None, language: str = "en") -> Document:
    """Create a styled document: page setup, masthead (title block), footer.

    `template` = school letterhead .docx. When present the letterhead owns the
    masthead and footer: the style's own title rules / two-tone bars and its
    SketchCast footer are DEMOTED (skipped) so they never collide with the
    school branding; section styling below stays.
    """
    style = STYLES[STYLE_BY_KIND.get(kind or "", "classroom")]
    branded = bool(template)
    if branded:
        doc = Document(template)
        _clear_body(doc)
    else:
        doc = Document()
        _setup_page(doc, style)
    lang = (language or "en").strip().lower()
    ctx = _Ctx(style=style, lang=lang, rtl=lang in _RTL_LANGS, branded=branded)
    ctx.width = _content_width_dxa(doc)
    setattr(doc, _CTX_ATTR, ctx)

    # ── masthead ──
    if style.name == "exam_board" and not branded:
        p = _p(doc, ctx, align=WD_ALIGN_PARAGRAPH.CENTER, after=6, style="Title")
        _p_border(p, "top", sz=8, space=10)
        _p_border(p, "bottom", sz=4, space=8)
        _run(p, title, ctx, size=15, font=style.head_font, caps=True, spacing=2.2,
             color=style.ink)
        if subtitle:
            m = _p(doc, ctx, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=4)
            _run(m, subtitle, ctx, size=10.5, color=style.meta_color)
    elif style.name == "workbook" and not branded:
        p = _p(doc, ctx, after=7, style="Title")
        _p_border(p, "bottom", val="nil")  # Title style's built-in rule off
        _run(p, title, ctx, size=25, bold=True, color=style.ink)
        _two_tone_rule(doc, ctx, main_w=2230, second_w=940)
        if subtitle:
            m = _p(doc, ctx, before=10, after=12)
            _run(m, subtitle, ctx, size=10, color=style.meta_color)
    elif style.name == "classroom" and not branded:
        p = _p(doc, ctx, after=6, style="Title")
        _p_border(p, "bottom", sz=28, space=8, color=style.accents[0][0])  # thick teal rule
        _run(p, title, ctx, size=26, font=style.head_font, color=style.ink)
        if subtitle:
            m = _p(doc, ctx, before=6, after=8)
            _run(m, subtitle, ctx, size=9, color=style.meta_color)
    else:
        # Branded letterhead: demoted masthead — plain title + subtitle, no rules.
        p = _p(doc, ctx, after=4, style="Title")
        _p_border(p, "bottom", val="nil")
        _run(p, title, ctx, size=20, bold=True, color=style.ink)
        if subtitle:
            m = _p(doc, ctx, before=2, after=6)
            _run(m, subtitle, ctx, size=10, italic=True, color=style.meta_color)

    if not branded:
        _footer(doc, ctx, f"SketchCast · {title}")
    return doc


# ── Semantic building blocks ─────────────────────────────────────────────────

def heading(doc: Document, text: str, level: int = 1) -> None:
    """Section heading, dispatched on the document's style. The heading TEXT is
    the caller's exact (already-localized) string — never altered here."""
    ctx = _ctx(doc)
    st = ctx.style
    if st.name == "exam_board":
        if level == 0:
            # centered rule-bound block (e.g. the answer key's own page)
            p = _p(doc, ctx, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=10,
                   style="Title")
            _p_border(p, "top", sz=8, space=10)
            _p_border(p, "bottom", sz=4, space=8)
            _run(p, text, ctx, size=14, font=st.head_font, caps=True, spacing=2.2,
                 color=st.ink)
        else:
            p = _p(doc, ctx, before=(14 if level == 1 else 12), after=8,
                   style=f"Heading {min(level, 2)}")
            # A section head must never strand at a page bottom with its
            # questions/table starting overleaf (seen on exam Section D).
            p.paragraph_format.keep_with_next = True
            if level == 1 and not ctx.rtl:
                # small-caps head with a trailing hairline to the margin
                p.paragraph_format.tab_stops.add_tab_stop(
                    Cm(ctx.width / 567.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.LINES)
                _run(p, text, ctx, size=11.5, bold=True, small_caps=True, color=st.ink)
                _run(p, " \t", ctx, size=11.5, color=st.ink)
            else:
                r = _run(p, text, ctx, size=11.5 if level == 1 else 11, bold=True,
                         small_caps=not ctx.rtl, color=st.ink)
                if level == 1 and ctx.rtl:
                    _p_border(p, "bottom", sz=4, space=4)
    elif st.name == "workbook":
        if level == 0:
            p = _p(doc, ctx, before=4, after=7, style="Title")
            _p_border(p, "bottom", val="nil")
            _run(p, text, ctx, size=20, bold=True, color=st.ink)
            _two_tone_rule(doc, ctx, main_w=1510, second_w=720, height=70)
            _spacer(doc, 10)
            ctx.accent_i = -1  # answer panels restart the accent cycle
        else:
            ctx.accent_i += 1
            fill, on_fill, _dark, _tint = ctx.accent
            _spacer(doc, 2)
            _bar(doc, ctx, text, fill, on_fill, full_width=False,
                 heading_style=f"Heading {min(level, 2)}", size=11.5)
    else:  # classroom
        if level == 0:
            _spacer(doc, 2)
            _bar(doc, ctx, text, "14181F", "FFFFFF", full_width=True,
                 heading_style="Title", size=12)
            _spacer(doc, 6)
            ctx.accent_i = -1
        elif level == 1:
            ctx.accent_i += 1
            _spacer(doc, 4)
            _bar(doc, ctx, text, st.accents[0][0], "FFFFFF", full_width=True,
                 heading_style="Heading 1", size=12)
        else:
            p = _p(doc, ctx, before=10, after=4, style=f"Heading {min(level, 2)}")
            _run(p, text, ctx, size=11.5, bold=True, color=st.accents[0][2])


def para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    ctx = _ctx(doc)
    p = _p(doc, ctx, after=6, line=1.15)
    _run(p, text, ctx, bold=bold, italic=italic)


def instructions(doc: Document, text: str) -> None:
    """Student instructions. Workbook: tinted callout with an accent side bar;
    other styles: an italic lead paragraph."""
    ctx = _ctx(doc)
    st = ctx.style
    if st.name == "workbook":
        fill, _on, dark, tint = st.accents[0]
        tbl = doc.add_table(rows=1, cols=1)
        _table_borders(tbl, none=True)
        _table_fixed(tbl, [ctx.width])
        _table_bidi(tbl, ctx)
        cell = tbl.rows[0].cells[0]
        side = ("single", 28, fill)
        # Cell borders swap under w:bidiVisual too — "left" IS the reading
        # edge in both directions (renders right for RTL, verified).
        _cell_borders(cell, top=None, bottom=None, left=side, right=None)
        _shade_cell(cell, tint)
        _cell_margins(cell, top=150, left=240, bottom=150, right=200)
        p = _fmt_cell_par(cell, ctx)
        p.paragraph_format.line_spacing = 1.15
        _run(p, text, ctx, size=10.5, italic=True, color="1E655E")
        _spacer(doc, 8)
    else:
        p = _p(doc, ctx, before=8, after=8, line=1.15)
        _run(p, text, ctx, size=10.5, italic=True)


def labelled(doc: Document, label: str, value: str) -> None:
    ctx = _ctx(doc)
    p = _p(doc, ctx, after=4, line=1.15)
    # label follows the CURRENT section's accent (workbook cycles per chip)
    accent_dark = ctx.accent[2] if ctx.style.name != "exam_board" else ctx.style.ink
    _run(p, f"{label}: ", ctx, bold=True, color=accent_dark)
    _run(p, value, ctx)


def bullets(doc: Document, items: Iterable[str]) -> None:
    ctx = _ctx(doc)
    # A fabricated "List Bullet" would carry no numbering part — items would
    # render flush and markless. On a letterhead that lacks the style, a
    # literal glyph with a matching indent is the honest fallback (mirrors to
    # the right edge under bidi like every logical-start run).
    try:
        bullet_style = doc.styles["List Bullet"]
    except KeyError:
        bullet_style = None
    for it in items:
        if str(it).strip():
            if bullet_style is not None:
                p = doc.add_paragraph(str(it), style=bullet_style)
            else:
                p = doc.add_paragraph(f"• {it}")
                p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.name = ctx.style.body_font
                run.font.size = Pt(ctx.style.body_size)
            if ctx.rtl:
                _p_bidi(p)  # bidi alone: default start alignment = right


# One LLM-emitted question-number prefix: "1.", "12)", "Q3:", "Question 4]" …
_LEADING_NUM_RE = re.compile(r"^\s*(?:Q(?:uestion)?\s*)?\d{1,3}\s*[.)\]:]\s+", re.IGNORECASE)


def strip_leading_number(text: str) -> str:
    """Remove ONE LLM-emitted "1." / "Q2)" prefix from QUESTION text — the
    renderer owns numbering, so a model that self-numbers must not double it.

    Never apply to answers: an answer like "42." would vanish."""
    return _LEADING_NUM_RE.sub("", str(text), count=1)


def txt(value) -> str:
    """None-safe str() for answer-key cells: an explicit JSON null must render
    as numbered()'s em-dash placeholder, not the string "None". Preserves
    falsy-but-real answers ("0", False) — only None becomes blank."""
    return "" if value is None else str(value)


# Renderer-assembled question strings: leading "N. " and (exam style) a
# builder-appended trailing marks bracket separated by 2+ spaces, so a
# question that legitimately ends in brackets is never mistaken for marks.
_Q_NUM_RE = re.compile(r"^(\d{1,3})\.\s+(.*)$", re.S)
_Q_MARKS_RE = re.compile(r"\s{2,}(\[[^\[\]]+\])\s*$")

_Q_INDENT = Cm(0.75)


def question(doc: Document, text: str, *, first: bool = False) -> None:
    """One question paragraph. The caller assembles the exact string (number,
    text, optional trailing marks) — this only renders it per style:
    accent-colored bold number, hanging indent, and for the exam style the
    marks bracket right-tabbed to the margin (kept inline for RTL, where the
    physical right-tab would land on the leading edge)."""
    ctx = _ctx(doc)
    st = ctx.style
    p = _p(doc, ctx, before=(12 if first else 10) if st.name != "exam_board" else 0,
           after=2 if st.name != "exam_board" else 9,
           left=_Q_INDENT, hang=_Q_INDENT, line=1.15)
    # Tab-based gutter numbers + right-tabbed marks are LTR-only: tab stop
    # positions don't mirror under w:bidi, so RTL keeps everything inline
    # (the hanging indent itself mirrors — bidi swaps w:ind sides).
    exam_tabs = st.name == "exam_board" and not ctx.rtl
    if exam_tabs:
        p.paragraph_format.tab_stops.add_tab_stop(_Q_INDENT, WD_TAB_ALIGNMENT.LEFT)

    body = str(text)
    marks = None
    if exam_tabs:
        mm = _Q_MARKS_RE.search(body)
        if mm:
            marks = mm.group(1)
            body = body[: mm.start()]
    m = _Q_NUM_RE.match(body)
    num, rest = (m.group(1), m.group(2)) if m else (None, body)

    if st.name == "exam_board":
        if num is not None:
            _run(p, f"{num}.", ctx)
            _run(p, "\t" if exam_tabs else " ", ctx)
        _run(p, rest, ctx)
        if marks is not None:
            p.paragraph_format.tab_stops.add_tab_stop(
                Cm(ctx.width / 567.0), WD_TAB_ALIGNMENT.RIGHT)
            _run(p, "\t", ctx)
            _run(p, marks, ctx)
    else:
        accent_dark = ctx.accent[2] if st.name == "workbook" else st.body_color
        if num is not None:
            _run(p, f"{num}.", ctx, bold=True, color=accent_dark)
            _run(p, " " + rest, ctx)
        else:
            _run(p, rest, ctx)


def numbered(doc: Document, items: Iterable[str]) -> None:
    """Manually numbered list that restarts at 1 on EVERY call.

    Word's "List Number" style is one document-wide counter that never restarts,
    so Section B would continue Section A and the answer key would continue the
    paper. Blank items render as an em-dash placeholder instead of being skipped
    so paper item N and answer-key item N always stay aligned."""
    for n, it in enumerate(items, 1):
        question(doc, f"{n}. {str(it).strip() or '—'}", first=(n == 1))


def writing_lines(doc: Document, n: int) -> None:
    """n ruled writing lines for handwritten answers (dotted for the exam
    style, light solid otherwise).

    Word gotcha (mockups): consecutive paragraphs whose pBdr borders are
    IDENTICAL merge into one border group — only the last bottom rule draws.
    Alternating w:space (1, 2, 1, …) keeps every line its own rule."""
    ctx = _ctx(doc)
    st = ctx.style
    val = "dotted" if st.name == "exam_board" else "single"
    color = _DOT_GREY if st.name == "exam_board" else st.line_color
    for k in range(max(0, int(n))):
        p = _p(doc, ctx, before=24 if k == 0 else 20, after=0, left=_Q_INDENT)
        p.paragraph_format.right_indent = Cm(0.35)
        _p_border(p, "bottom", val=val, sz=6, space=1 + (k % 2), color=color)
        pPr = p._p.get_or_add_pPr()  # noqa: SLF001
        rPr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "16")  # 8pt paragraph mark keeps the rule tight
        rPr.append(sz)
        pPr.insert_element_before(rPr, "w:sectPr", "w:pPrChange")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Data table (match-the-columns, vocabulary, lesson flow …), styled per
    kind, direction-aware (w:bidiVisual mirrors columns for RTL)."""
    ctx = _ctx(doc)
    st = ctx.style
    t = doc.add_table(rows=1, cols=len(headers))
    t.autofit = True
    _table_bidi(t, ctx)
    if st.name == "exam_board":
        _table_borders(t, sz=4, color="000000")
        header_fill = None
        header_color = st.ink
    elif st.name == "workbook":
        _table_borders(t, sz=4, color=st.line_color)
        header_fill = ctx.accent[0]
        header_color = ctx.accent[1]
    else:
        _table_borders(t, sz=4, color=_HAIR)
        header_fill = "F0F3F5"
        header_color = "5A616B"
    # cell padding
    tblPr = t._tbl.tblPr  # noqa: SLF001
    mar = OxmlElement("w:tblCellMar")
    for side, w in (("top", 50), ("left", 110), ("bottom", 50), ("right", 110)):
        el = OxmlElement("w:" + side)
        el.set(qn("w:w"), str(w))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)

    for i, htext in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        if header_fill:
            _shade_cell(cell, header_fill)
        p = _fmt_cell_par(cell, ctx)
        _run(p, htext, ctx, size=10, bold=True, color=header_color,
             small_caps=(st.name == "exam_board" and not ctx.rtl))
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row[: len(headers)]):
            cells[i].text = ""
            p = _fmt_cell_par(cells[i], ctx)
            _run(p, str(val), ctx, size=10.5)


def marks_table(doc: Document, entries: list[tuple[str, int]]) -> None:
    """Examiner marks summary (exam style): Section / Marks rows + Total,
    floated to the trailing edge (right for LTR, left for RTL)."""
    ctx = _ctx(doc)
    lang = ctx.lang
    t = doc.add_table(rows=len(entries) + 2, cols=2)
    # RIGHT = trailing edge in both directions: table jc swaps under
    # w:bidiVisual, so this floats right for LTR and left for RTL (mirrored).
    t.alignment = WD_TABLE_ALIGNMENT.RIGHT
    _table_borders(t, sz=4, color="000000")
    _table_fixed(t, [1470, 1020])
    _table_bidi(t, ctx)
    tblPr = t._tbl.tblPr  # noqa: SLF001
    mar = OxmlElement("w:tblCellMar")
    for side, w in (("top", 40), ("left", 100), ("bottom", 40), ("right", 100)):
        el = OxmlElement("w:" + side)
        el.set(qn("w:w"), str(w))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)
    data = ([(_t("section", lang), _t("marks", lang))]
            + [(label, str(marks)) for label, marks in entries]
            + [(_t("total", lang), str(sum(m for _l, m in entries)))])
    for i, (c1, c2) in enumerate(data):
        row = t.rows[i]
        _row_height(row, 310)
        hdr = i == 0
        tot = i == len(data) - 1
        for j, cell_text in enumerate((c1, c2)):
            cell = row.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = _fmt_cell_par(cell, ctx, align=WD_ALIGN_PARAGRAPH.CENTER)
            _run(p, cell_text, ctx, size=9.5, bold=(hdr or tot),
                 small_caps=(hdr and not ctx.rtl))
    _spacer(doc, 6)


def tf_boxes(doc: Document, statements: Iterable[str]) -> None:
    """True/False with ruled answer boxes (exam style): each statement in a
    row with an empty boxed cell and a [1] mark, plus the localized cue line.
    Numbering restarts at 1 (same contract as numbered())."""
    ctx = _ctx(doc)
    items = [str(s).strip() or "—" for s in statements]
    if not items:
        return
    cue = _p(doc, ctx, after=6, left=_Q_INDENT)
    # The cue names the SAME localized True/False pair the answer key prints
    # (strings.tf_word) — one table entry serves both, so they can never drift.
    cue_text = _t("tf_box_cue", ctx.lang).format(
        t=tf_word(True, ctx.lang), f=tf_word(False, ctx.lang))
    _run(cue, cue_text, ctx, size=10, italic=True)

    rows = 2 * len(items) - 1  # spacer rows between box rows keep boxes apart
    t = doc.add_table(rows=rows, cols=3)
    _table_borders(t, none=True)
    text_w = ctx.width - 1700 - 700
    _table_fixed(t, [text_w, 1000, 700])
    _table_bidi(t, ctx)
    for k in range(1, rows, 2):  # spacer rows
        sp = t.rows[k]
        _row_height(sp, 170, rule="exact")
        p = _fmt_cell_par(sp.cells[0], ctx)
        r = p.add_run("")
        r.font.size = Pt(4)
    for n, statement in enumerate(items, 1):
        row = t.rows[(n - 1) * 2]
        _row_height(row, 480)
        c0, c1, c2 = row.cells
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p0 = _fmt_cell_par(c0, ctx)
        p0.paragraph_format.left_indent = _Q_INDENT
        p0.paragraph_format.first_line_indent = -_Q_INDENT
        if ctx.rtl:  # tabs don't mirror under bidi — inline number instead
            _run(p0, f"{n}. ", ctx)
        else:
            p0.paragraph_format.tab_stops.add_tab_stop(_Q_INDENT, WD_TAB_ALIGNMENT.LEFT)
            _run(p0, f"{n}.", ctx)
            _run(p0, "\t", ctx)
        _run(p0, statement, ctx)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _cell_borders(c1, top=("single", 6, "000000"), left=("single", 6, "000000"),
                      bottom=("single", 6, "000000"), right=("single", 6, "000000"))
        _fmt_cell_par(c1, ctx)
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # jc RIGHT = trailing edge on both sides (it swaps under w:bidi)
        p2 = _fmt_cell_par(c2, ctx, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _run(p2, "[1]", ctx)
    _spacer(doc, 4)


def answer_section(doc: Document, heading_text: str, items: Iterable[str]) -> None:
    """One answer-key section. Workbook: a tinted panel (accent side bar +
    tinted fill) with the heading and each numbered answer inside; other
    styles: a level-2 heading + numbered list. Numbering contract identical
    to numbered() — restarts at 1, em-dash placeholder for blanks."""
    ctx = _ctx(doc)
    st = ctx.style
    if st.name != "workbook":
        heading(doc, heading_text, 2)
        numbered(doc, items)
        return
    ctx.accent_i += 1
    fill, _on, dark, tint = ctx.accent
    tbl = doc.add_table(rows=1, cols=1)
    _table_borders(tbl, none=True)
    _table_fixed(tbl, [ctx.width])
    _table_bidi(tbl, ctx)
    cell = tbl.rows[0].cells[0]
    side = ("single", 24, fill)
    # "left" border = reading edge under w:bidiVisual (see instructions())
    _cell_borders(cell, top=None, bottom=None, left=side, right=None)
    _shade_cell(cell, tint)
    _cell_margins(cell, top=140, left=240, bottom=150, right=200)
    ph = _fmt_cell_par(cell, ctx)
    ph.style = _ensure_style(doc, "Heading 2")
    ph.paragraph_format.space_before = Pt(0)
    ph.paragraph_format.space_after = Pt(4)
    if ctx.rtl:
        _p_bidi(ph)
    _run(ph, heading_text, ctx, size=11, bold=True, color=dark)
    for n, it in enumerate(items, 1):
        pa = _cell_extra_par(cell, ctx)
        pa.paragraph_format.line_spacing = 1.2
        _run(pa, f"{n}.", ctx, bold=True, color=dark)
        _run(pa, " " + (str(it).strip() or "—"), ctx)
    _spacer(doc, 8)


def end_of_paper(doc: Document) -> None:
    """Centered end-of-paper register mark (exam style), localized."""
    ctx = _ctx(doc)
    p = _p(doc, ctx, align=WD_ALIGN_PARAGRAPH.CENTER, before=30, after=0)
    _run(p, f"— {_t('end_of_paper', ctx.lang)} —", ctx, size=9.5, spacing=1.2)


def save(doc: Document, path: str | Path) -> Path:
    path = Path(path)
    doc.save(str(path))
    return path


# ── Prompt grounding (unchanged) ─────────────────────────────────────────────

def chapter_excerpt(chapter: dict, limit: int = 6000) -> str:
    """Flatten a chapter's section text into a single excerpt for prompting."""
    parts: list[str] = []
    for sec in chapter.get("sections", []) or []:
        title = (sec.get("section_title") or "").strip()
        content = (sec.get("content") or "").strip()
        if title:
            parts.append(title)
        if content:
            parts.append(content)
    text = "\n".join(parts).strip()
    return text[:limit] if text else (chapter.get("title") or "")


def concept_names(analysis: dict, limit: int = 12) -> list[str]:
    """Best-effort concept names from the Agent 2 analysis (shape-tolerant)."""
    names: list[str] = []
    for key in ("concepts", "key_concepts"):
        for c in analysis.get(key, []) or []:
            name = c.get("name") or c.get("concept") or c.get("term") if isinstance(c, dict) else str(c)
            if name:
                names.append(str(name))
    return names[:limit]


def chapter_grounding(book: dict, chapter: dict, analysis: dict) -> str:
    """The shared grounding block prepended to EVERY artifact prompt for a chapter.

    Must be byte-identical across all of a book's artifacts (worksheet, exam, …)
    so Claude prompt-caches it — written once, re-read at ~0.1x on the rest. Keep
    it deterministic: only book/chapter/analysis-derived text, no per-artifact
    params, no timestamps. Artifact-specific instructions go AFTER this, in each
    builder's own prompt, and reference "the chapter above"."""
    grade = book.get("grade") or "school"
    subject = book.get("subject") or "general"
    title = chapter.get("title") or "Chapter"
    concepts = ", ".join(concept_names(analysis)) or "(see content)"
    return (
        f"Audience: {grade} {subject} students.\n"
        f'Chapter: "{title}"\n'
        f"Key concepts: {concepts}\n\n"
        f'Chapter content:\n"""\n{chapter_excerpt(chapter)}\n"""'
    )
