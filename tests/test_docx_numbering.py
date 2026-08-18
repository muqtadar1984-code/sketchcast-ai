"""Renderer-owned numbering across the docx builders — every section restarts
at 1 and the answer key carries EXACTLY the paper's numbers. Regression for the
Word "List Number" style: ONE document-wide counter that never restarts, so
Section B continued Section A and the Answer Key continued the paper.
(The worksheet builder's numbering test lives in tests/test_worksheet.py.)
"""

from __future__ import annotations

import re

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from docgen import activity, case_study, exam, exam_paper


class _StubClient:
    """Returns a fixed payload (no network)."""

    def __init__(self, data):
        self._data = data

    def analyze(self, prompt, max_tokens=0, **k):
        return {"data": self._data}


_BOOK = {"grade": "Grade 7", "subject": "Science"}
_CHAPTER = {"title": "Cells", "sections": [{"section_title": "Cells", "content": "A cell is the basic unit."}]}


def _iter_paras(doc):
    """Every paragraph in document order — the style layer renders section
    headings/chips and T/F rows inside single-cell tables, so table-cell
    paragraphs count too."""
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            for row in Table(child, doc).rows:
                for cell in row.cells:
                    yield from cell.paragraphs


def _text(p):
    """Whitespace-normalized paragraph text (the exam style tabs its question
    numbers and right-aligns marks with a tab — layout, not content)."""
    return re.sub(r"\s+", " ", p.text).strip()


def _sections(doc):
    """Ordered (heading text, [non-empty paragraph texts under it]) pairs."""
    out = []
    for p in _iter_paras(doc):
        if (p.style.name or "").startswith(("Heading", "Title")):
            out.append((_text(p), []))
        elif out and _text(p):
            out[-1][1].append(_text(p))
    return out


def _numbered(texts):
    """Just the renderer-numbered items of a section — skips style chrome the
    exam style adds around them ("[1]" mark cells, the write-in-box cue)."""
    return [t for t in texts if re.match(r"\d+\. ", t)]


def _no_list_number(doc):
    assert all(p.style.name != "List Number" for p in doc.paragraphs)


# ── activity.py — each activity's steps restart at 1 ────────────────────────────

_ACTIVITY_PAYLOAD = {
    "intro": "Two activities.",
    "activities": [
        {"name": "Model a cell", "objective": "o", "grouping": "pairs", "duration_minutes": 10,
         "materials": ["clay"], "steps": ["Roll the clay", "Label the parts"],
         "teacher_facilitation": "Walk around", "success_looks_like": "Labels correct"},
        {"name": "Cell walk", "objective": "o", "grouping": "whole class", "duration_minutes": 15,
         "materials": ["cards"], "steps": ["Hand out cards", "Students find partners", "Discuss"],
         "teacher_facilitation": "Moderate", "success_looks_like": "Pairs correct"},
    ],
}


def test_activity_steps_restart_at_one_per_activity(tmp_path):
    # DELIBERATE pin update (2026-08-18 student/teacher split): build() now
    # returns [student_sheet, teacher_notes]; the steps live in the student sheet.
    out, _notes = activity.build(_BOOK, _CHAPTER, {}, _StubClient(_ACTIVITY_PAYLOAD),
                                 {"num_activities": 2}, tmp_path)
    doc = Document(str(out))
    _no_list_number(doc)
    texts = [_text(p) for p in _iter_paras(doc) if _text(p)]
    # Activity 2's steps must restart at 1, not continue Activity 1's count at 3.
    second_steps = texts.index("Steps:", texts.index("Steps:") + 1)
    assert texts[second_steps + 1] == "1. Hand out cards"
    assert texts[second_steps + 2] == "2. Students find partners"
    assert texts[second_steps + 3] == "3. Discuss"


# ── exam_paper.py — separate answer key mirrors the paper's numbers ─────────────
# DELIBERATE pin update (2026-08-18 student/teacher split): the answer key is no
# longer an inline page — build() returns [paper, key] like exam.py, and the
# numbering contract (key item N answers paper item N, every section restarts
# at 1) now holds ACROSS the two files.

_EXAM_PAPER_PAYLOAD = {
    "title": "Cells Test",
    "instructions": "Answer everything.",
    "fill_blank": [
        {"q": "1. The ____ controls the cell.", "answer": "nucleus"},
        {"q": "Plant cells have a ____ wall.", "answer": "cell"},
    ],
    "true_false": [
        {"statement": "Animal cells have chloroplasts.", "answer": False},
        {"statement": "The nucleus stores DNA.", "answer": True},
    ],
    "match_column": [
        {"left": "Nucleus", "right": "Controls the cell"},
        {"left": "Vacuole", "right": "Stores water"},
    ],
    "subjective": [{"q": "Q1) Explain osmosis.", "marks": 5, "answer_outline": "Movement of water."}],
}


def test_exam_paper_key_numbers_match_paper(tmp_path):
    paths = exam_paper.build(_BOOK, _CHAPTER, {}, _StubClient(_EXAM_PAPER_PAYLOAD),
                             {"objective": {"fill_blank": 2, "true_false": 2, "match_column": 2},
                              "subjective": 1}, tmp_path)
    paper_doc, key_doc = Document(str(paths[0])), Document(str(paths[1]))
    _no_list_number(paper_doc)
    _no_list_number(key_doc)
    paper = dict(_sections(paper_doc))
    key = dict(_sections(key_doc))
    # The student paper carries no answer-key page at all.
    assert "Answer Key" not in paper

    # Every section restarts at 1; LLM-emitted "1." / "Q1)" prefixes stripped once.
    assert paper["Section A — Fill in the blanks"][0] == "1. The ____ controls the cell."
    assert _numbered(paper["Section B — True or False"])[0].startswith("1. ")
    assert _numbered(paper["Section D — Long answer"])[0] == "1. Explain osmosis. [5 marks]"

    # The key carries exactly the paper's numbers, section by section.
    for name in ("Section A — Fill in the blanks", "Section B — True or False"):
        assert [t.split(".")[0] for t in key[name]] == \
            [t.split(".")[0] for t in _numbered(paper[name])], name
    assert key["Section B — True or False"] == ["1. False", "2. True"]
    assert key["Section D — Long answer"] == ["1. Movement of water."]
    # Match key is "N. <letter>", not the old self-numbered "1. 1 → C".
    assert all(re.fullmatch(r"\d+\. [A-Z]", t) for t in key["Section C — Match the columns"])


def test_exam_paper_stray_strings_cannot_skew_key_alignment(tmp_path):
    """A non-dict item in the model's list must not consume a paper number.

    Regression: the long-answer loop skipped strays INSIDE enumerate while the
    key filtered them OUT — paper read 1./3. against a key of 1./2. Non-dicts
    are now dropped once at extraction, so both sides enumerate the same list.
    A null answer renders as the em-dash slot, never the string "None"."""
    payload = {
        "title": "Cells Test",
        "fill_blank": [{"q": "The ____ controls the cell.", "answer": None}],
        "subjective": [
            {"q": "First real question.", "marks": 2, "answer_outline": "one"},
            "stray string the model emitted",
            {"q": "Second real question.", "marks": 3, "answer_outline": "two"},
        ],
    }
    paths = exam_paper.build(_BOOK, _CHAPTER, {}, _StubClient(payload),
                             {"objective": {"fill_blank": 1}, "subjective": 3}, tmp_path)
    paper = dict(_sections(Document(str(paths[0]))))
    key = dict(_sections(Document(str(paths[1]))))

    subj = [t for t in paper["Section B — Long answer"] if re.match(r"\d+\.", t)]
    assert subj[0].startswith("1. First real question.")
    assert subj[1].startswith("2. Second real question.")  # not "3."
    assert key["Section B — Long answer"] == ["1. one", "2. two"]
    assert key["Section A — Fill in the blanks"] == ["1. —"]  # null answer, not "None"


# ── case_study.py — guidance N answers question N ───────────────────────────────

def test_case_study_guidance_numbers_match_questions(tmp_path):
    """Same stray-string regression as the test paper: questions and teacher
    notes must enumerate the same filtered list.

    DELIBERATE pin update (2026-08-18 student/teacher split): the teacher notes
    are now a SEPARATE document — guidance N still answers question N, but
    across the two files."""
    payload = {
        "title": "Case",
        "scenario": "A scenario.",
        "discussion_questions": [
            {"q": "1. Why did the cell divide?", "guidance": "growth"},
            "stray string",
            {"q": "What limits cell size?", "guidance": "surface area"},
        ],
    }
    paths = case_study.build(_BOOK, _CHAPTER, {}, _StubClient(payload),
                             {"num_questions": 3}, tmp_path)
    secs = dict(_sections(Document(str(paths[0]))))
    qs = [t for t in secs["Discussion questions"] if re.match(r"\d+\.", t)]
    assert qs == ["1. Why did the cell divide?", "2. What limits cell size?"]
    # The student document carries no teacher notes; the separate notes doc does.
    assert "Teacher notes — answer guidance" not in secs
    key_secs = dict(_sections(Document(str(paths[1]))))
    assert key_secs["Teacher notes — answer guidance"] == ["1. growth", "2. surface area"]


# ── exam.py — cumulative exam: separate key doc restarts per section ────────────

_EXAM_PAYLOAD = {
    "title": "Term Exam",
    "instructions": "Answer everything.",
    "mcq": [
        {"q": "1. Which organelle makes energy?",
         "options": ["Nucleus", "Mitochondria", "Vacuole", "Wall"], "answer": "B", "marks": 1},
        {"q": "Which organelle stores water?",
         "options": ["Nucleus", "Mitochondria", "Vacuole", "Wall"], "answer": "C", "marks": 1},
    ],
    "fill_blank": [{"q": "The ____ controls the cell.", "answer": "nucleus", "marks": 1}],
    "true_false": [{"statement": "Cells are alive.", "answer": True, "marks": 1}],
    "match_column": [
        {"left": "Nucleus", "right": "Controls the cell"},
        {"left": "Vacuole", "right": "Stores water"},
    ],
    "short_answer": [{"q": "2. Define diffusion.", "answer": "Movement of particles.", "marks": 2}],
    "long_answer": [{"q": "Describe the cell cycle.", "answer_outline": "Phases in order.", "marks": 5}],
}


def test_cumulative_exam_key_restarts_per_section(tmp_path):
    paths = exam.build(_BOOK, _CHAPTER, {}, _StubClient(_EXAM_PAYLOAD),
                       {"counts": {"mcq": 2, "fill_blank": 1, "true_false": 1, "match_column": 2,
                                   "short_answer": 1, "long_answer": 1}}, tmp_path)
    paper, key = Document(str(paths[0])), Document(str(paths[1]))
    _no_list_number(paper)
    _no_list_number(key)

    paper_secs = dict(_sections(paper))
    # LLM-emitted prefixes stripped on both the manual and numbered() paths.
    assert paper_secs["Section A — Multiple choice"][0] == "1. Which organelle makes energy? [1]"
    assert paper_secs["Section B — Fill in the blanks"] == ["1. The ____ controls the cell."]
    assert _numbered(paper_secs["Section E — Short answer"])[0] == "1. Define diffusion. [2]"

    key_secs = dict(_sections(key))
    # Each key section restarts at 1 (True/False used to continue at 4 after
    # 2 MCQ + 1 fill) and mirrors the paper's numbers.
    assert key_secs["Section A — Multiple choice"] == ["1. B", "2. C"]
    assert key_secs["Section B — Fill in the blanks"] == ["1. nucleus"]
    assert key_secs["Section C — True or False"] == ["1. True"]
    assert key_secs["Section E — Short answer"] == ["1. Movement of particles."]
    assert key_secs["Section F — Long answer"] == ["1. Phases in order."]
    # Match key is "N. <letter>", not the old self-numbered "1. 1 → C".
    assert all(re.fullmatch(r"\d+\. [A-Z]", t) for t in key_secs["Section D — Match the columns"])
