"""Regression test for the 2026-08-18 production failure: letterheads without
Word's built-in styles.

Branded documents build on a teacher's uploaded letterhead .docx
(new_doc(template=...)). Word keeps built-in styles LATENT until a document
actually uses them, and python-docx cannot see latent styles — so a normal
letterhead defines no 'Title'/'Heading N'/'List Bullet', and the style layer's
doc.styles['Title'] lookup killed every document generation for that teacher
with "no style with name 'Title'" (the video, which writes no docx, survived —
5 of a kit's 6 pieces failed).

The fix (_ensure_style) fabricates a missing named style — based on Normal,
outline level for Heading N — and bullets() falls back to a literal glyph
rather than a numbering-less fabricated 'List Bullet'. These tests build a
letterhead stripped of every style the layer references and drive the full
branded surface for all six kinds; any KeyError regression fails loudly.

Run: python -m pytest tests/test_docx_letterhead_styles.py -q
"""

import io

import pytest
from docx import Document

from docgen import docx_builder as dx

STRIPPED = ("Title", "Heading 1", "Heading 2", "List Bullet")

ALL_KINDS = ("exam_paper", "exam", "worksheet", "activity", "lesson_plan", "case_study")


def _letterhead_without_builtin_styles() -> io.BytesIO:
    """A realistic letterhead: valid docx, Normal present, built-ins absent."""
    doc = Document()
    doc.add_paragraph("Sunrise Academy — letterhead")
    styles_el = doc.styles.element
    for name in STRIPPED:
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        styles_el.remove(st.element)
    # Prove the template reproduces the production precondition.
    for name in STRIPPED:
        with pytest.raises(KeyError):
            doc.styles[name]  # noqa: B018 — the lookup itself is the assertion
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@pytest.mark.parametrize("kind", ALL_KINDS)
def test_branded_doc_survives_styleless_letterhead(kind):
    doc = dx.new_doc("Chapter test", "Grade 7 · Science", template=_letterhead_without_builtin_styles(), kind=kind)
    # The full styled surface the builders drive: mastheads and section heads
    # are where the named-style lookups live.
    dx.heading(doc, "Section one", level=0)
    dx.heading(doc, "Sub-section", level=1)
    dx.heading(doc, "Sub-sub-section", level=2)
    dx.bullets(doc, ["First point", "Second point"])
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    reopened = Document(out)
    # The fabricated styles are real, named definitions in the saved file.
    assert reopened.styles["Title"] is not None
    assert reopened.styles["Heading 1"] is not None
    # Bullet fallback: items keep a visible mark even without the style.
    texts = [p.text for p in reopened.paragraphs]
    assert any(t.startswith("• First point") for t in texts)


def test_fabricated_heading_carries_outline_level():
    # Fabrication is lazy — the style appears at its first use, so drive a
    # section head onto the page before asserting on it.
    doc = dx.new_doc("T", template=_letterhead_without_builtin_styles(), kind="lesson_plan")
    dx.heading(doc, "Section", level=1)
    st = doc.styles["Heading 1"]
    lvls = st.element.findall(".//" + dx.qn("w:outlineLvl"))
    assert [e.get(dx.qn("w:val")) for e in lvls] == ["0"]


def test_unbranded_default_template_unchanged():
    # The default template already defines everything — _ensure_style must
    # return the existing styles, never duplicate them.
    doc = dx.new_doc("T", kind="worksheet")
    before = len(doc.styles.element)
    dx.heading(doc, "H", level=0)
    dx.bullets(doc, ["a"])
    assert len(doc.styles.element) == before
