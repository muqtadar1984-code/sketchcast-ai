"""Document-builder i18n — every user-visible framework string the builders
emit (section headings, Answer Key, Column A/B, True/False key values, marks
brackets, labels, default titles) is localized per lesson language, and the
consistency rules hold WITHIN a document:

· Arabic-script docs (ar, ms-arab) enumerate sections / match options / MCQ
  options with abjad letters (أ ب ج …) — and the answer key's letters come
  from the SAME sequence, or the key can't be graded against the paper.
· The key's True/False values are the same localized pair the paper's cue
  names (ar صواب/خطأ — NEVER English True/False).
· ms-arab is Jawi: Arabic SCRIPT but MALAY vocabulary (بهاݢين, بنر/ڤلسو) —
  never the Arabic-language terms (القسم, صواب/خطأ).
· English output is byte-identical to the pre-i18n literals (the pinned tests
  in test_worksheet.py / test_docx_numbering.py stay untouched; the exact-
  string asserts here would catch en drift in the _STRINGS table).
"""

from __future__ import annotations

import re

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from docgen import exam, exam_paper, worksheet
from docgen.strings import letters


class _StubClient:
    """Returns a fixed payload (no network)."""

    def __init__(self, data):
        self._data = data

    def analyze(self, prompt, max_tokens=0, **k):
        return {"data": self._data}


def _iter_paras(doc):
    """Every paragraph in document order, table-cell paragraphs included (the
    style layer renders chips/panels/T-F rows inside tables)."""
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            for row in Table(child, doc).rows:
                for cell in row.cells:
                    yield from cell.paragraphs


def _text(p):
    return re.sub(r"\s+", " ", p.text).strip()


def _blob(doc):
    return "\n".join(_text(p) for p in _iter_paras(doc) if _text(p))


def _sections(doc):
    out = []
    for p in _iter_paras(doc):
        if (p.style.name or "").startswith(("Heading", "Title")):
            out.append((_text(p), []))
        elif out and _text(p):
            out[-1][1].append(_text(p))
    return out


def _match_table(doc, headers):
    found = [t for t in doc.tables
             if [c.text for c in t.rows[0].cells] == headers]
    assert len(found) == 1, f"match table with headers {headers} not found"
    return found[0]


_ABJAD_FIRST = ["أ", "ب", "ج", "د", "هـ", "و"]

_WS_PAYLOAD_AR = {
    "title": "ورقة عمل: الخلية",
    "instructions": "أجب عن جميع الأسئلة.",
    "fill_blank": [
        {"q": "تتحكم ____ في أنشطة الخلية.", "answer": "النواة"},
        {"q": "تمتلك الخلايا النباتية جدارًا ____.", "answer": "خلويًا"},
    ],
    "true_false": [
        {"statement": "تحتوي الخلايا الحيوانية على بلاستيدات خضراء.", "answer": False},
        {"statement": "تخزن النواة المادة الوراثية.", "answer": True},
    ],
    "match_column": [
        {"left": "النواة", "right": "تتحكم في الخلية"},
        {"left": "البلاستيدة الخضراء", "right": "تلتقط ضوء الشمس"},
        {"left": "الفجوة", "right": "تخزن الماء"},
    ],
    "short_answer": [{"q": "لماذا تكون الخلايا النباتية صلبة؟", "answer": "بسبب الجدار الخلوي.", "work_space_lines": 1}],
}

_BOOK_AR = {"grade": "الصف السابع", "subject": "العلوم"}
_CHAP_AR = {"title": "الخلية", "sections": [{"section_title": "الخلية", "content": "الخلية وحدة البناء."}]}


def _build_ws(payload, lang, tmp_path, book=None, chap=None):
    out = worksheet.build(
        book=book or _BOOK_AR, chapter=chap or _CHAP_AR, analysis={},
        client=_StubClient(payload),
        params={"num_questions": 10, "include_answer_key": True},
        out_dir=tmp_path, language=lang,
    )
    return Document(str(out))


# ── Arabic ───────────────────────────────────────────────────────────────────

def test_arabic_worksheet_headings_are_native_with_abjad_letters(tmp_path):
    doc = _build_ws(_WS_PAYLOAD_AR, "ar", tmp_path)
    blob = _blob(doc)
    assert "القسم أ — املأ الفراغات" in blob
    assert "القسم ب — صواب أو خطأ" in blob
    assert "القسم ج — صل بين العمودين" in blob
    assert "القسم د — الإجابة القصيرة" in blob
    assert "نموذج الإجابة" in blob
    # No English framework strings anywhere in the document body.
    assert not re.search(r"\b(Section|Answer Key|Column|True|False|Difficulty|marks)\b", blob)


def test_arabic_key_true_false_uses_the_cue_pair_not_english(tmp_path):
    doc = _build_ws(_WS_PAYLOAD_AR, "ar", tmp_path)
    # Paper and key share the heading string; dict(_sections) keeps the LAST
    # occurrence — the answer-key panel, whose items are the key values.
    secs = dict(_sections(doc))
    # Answers were (False, True) → the key prints the localized pair.
    assert secs["القسم ب — صواب أو خطأ"] == ["1. خطأ", "2. صواب"]


def test_arabic_match_letters_shared_between_table_and_key(tmp_path):
    doc = _build_ws(_WS_PAYLOAD_AR, "ar", tmp_path)
    t = _match_table(doc, ["العمود أ", "العمود ب"])
    right_by_letter = {}
    for row in t.rows[1:]:
        letter, right = row.cells[1].text.split(". ", 1)
        assert letter in _ABJAD_FIRST  # abjad, not A/B/C
        right_by_letter[letter] = right
    secs = dict(_sections(doc))
    key_lines = secs["القسم ج — صل بين العمودين"]
    key_lines = [t for t in key_lines if re.match(r"\d+\. ", t)]
    assert len(key_lines) == 3
    for i, line in enumerate(key_lines):
        n, letter = line.split(". ", 1)
        assert int(n) == i + 1
        assert letter in right_by_letter, f"key letter {letter!r} not in the paper's table"
        assert right_by_letter[letter] == _WS_PAYLOAD_AR["match_column"][i]["right"]


def test_arabic_exam_paper_marks_bracket_and_letters(tmp_path):
    payload = {
        "title": "ورقة اختبار",
        "instructions": "أجب عن جميع الأسئلة.",
        "fill_blank": [{"q": "تتحكم ____ في الخلية.", "answer": "النواة"}],
        "true_false": [{"statement": "الخلايا حية.", "answer": True}],
        "subjective": [{"q": "اشرح الخاصية الأسموزية.", "marks": 5, "answer_outline": "انتقال الماء."}],
    }
    out = exam_paper.build(_BOOK_AR, _CHAP_AR, {}, _StubClient(payload),
                           {"objective": {"fill_blank": 1, "true_false": 1},
                            "subjective": 1, "include_answer_key": True},
                           tmp_path, language="ar")
    doc = Document(str(out))
    blob = _blob(doc)
    assert "[5 درجات]" in blob            # localized marks wording
    assert "[5 marks]" not in blob
    assert "اكتب صواب أو خطأ في كل مربع." in blob  # T/F cue = the key's pair
    assert "القسم أ — املأ الفراغات" in blob
    assert "نهاية الورقة" in blob
    assert "نموذج الإجابة" in blob
    assert not re.search(r"\b(Section|Answer|Column|True|False|marks|Total)\b", blob)
    # Marks table letters are abjad too (أ = section A row).
    marks_tbl = doc.tables[0]
    first_col = [r.cells[0].text for r in marks_tbl.rows]
    assert first_col[0] == "القسم" and first_col[-1] == "المجموع"
    assert set(first_col[1:-1]) <= set(_ABJAD_FIRST)


def test_arabic_cumulative_exam_mcq_key_letters_match_the_options(tmp_path):
    payload = {
        "title": "امتحان",
        "mcq": [
            {"q": "أي عضية تنتج الطاقة؟",
             "options": ["النواة", "الميتوكندريون", "الفجوة", "الجدار"], "answer": "B", "marks": 1},
            {"q": "أي تركيب نباتي فقط؟",
             "options": ["الغشاء", "السيتوبلازم", "الجدار الخلوي", "النواة"], "answer": "C", "marks": 1},
        ],
        "true_false": [{"statement": "الخلايا حية.", "answer": False}],
    }
    paths = exam.build(_BOOK_AR, _CHAP_AR, {}, _StubClient(payload),
                       {"counts": {"mcq": 2, "true_false": 1}}, tmp_path, language="ar")
    paper, key = Document(str(paths[0])), Document(str(paths[1]))
    paper_blob = _blob(paper)
    # Options are lettered with the abjad sequence…
    assert "أ) النواة" in paper_blob
    assert "ب) الميتوكندريون" in paper_blob
    # …and the key's letters come from the SAME sequence (B→ب, C→ج).
    key_secs = dict(_sections(key))
    assert key_secs["القسم أ — اختيار من متعدد"] == ["1. ب", "2. ج"]
    assert key_secs["القسم ب — صواب أو خطأ"] == ["1. خطأ"]
    assert "للمعلم فقط — لا يوزع على الطلاب." in _blob(key)
    assert not re.search(r"\b(Section|Answer|True|False)\b", _blob(key))


# ── Jawi (ms-arab): Arabic SCRIPT, MALAY vocabulary ──────────────────────────

def test_jawi_worksheet_is_malay_in_arabic_script_not_arabic(tmp_path):
    doc = _build_ws(_WS_PAYLOAD_AR, "ms-arab", tmp_path,
                    book={"grade": "تيڠکتن 1", "subject": "ساءينس"},
                    chap={"title": "سيل", "sections": []})
    blob = _blob(doc)
    # Jawi/Malay framework vocabulary…
    assert "بهاݢين أ — ايسي تمڤت کوسوڠ" in blob   # Bahagian A — Isi tempat kosong
    assert "بهاݢين ب — بنر اتاو ڤلسو" in blob     # Benar atau Palsu
    assert "سکيما جواڤن" in blob                  # Skema Jawapan
    # …never the Arabic-language terms for the same keys.
    assert "القسم" not in blob
    assert "صواب" not in blob and "خطأ" not in blob
    assert "نموذج الإجابة" not in blob
    assert "العمود" not in blob
    # Key True/False values are the Jawi pair.
    secs = dict(_sections(doc))
    assert secs["بهاݢين ب — بنر اتاو ڤلسو"] == ["1. ڤلسو", "2. بنر"]
    # Match letters still abjad (Arabic-script doc), shared with the key.
    t = _match_table(doc, ["لاجور أ", "لاجور ب"])
    letters_in_table = {row.cells[1].text.split(". ", 1)[0] for row in t.rows[1:]}
    assert letters_in_table <= set(_ABJAD_FIRST)
    key_lines = [x for x in secs["بهاݢين ج — ڤادنکن لاجور"] if re.match(r"\d+\. ", x)]
    assert {line.split(". ", 1)[1] for line in key_lines} == letters_in_table
    assert not re.search(r"\b(Section|Answer|Column|True|False)\b", blob)


# ── French: Latin script keeps A B C letters ─────────────────────────────────

def test_french_worksheet_headings_and_key(tmp_path):
    payload = {
        "title": "Fiche — La cellule",
        "instructions": "Répondez à toutes les questions.",
        "fill_blank": [{"q": "Le ____ contrôle la cellule.", "answer": "noyau"}],
        "true_false": [{"statement": "Les cellules animales ont des chloroplastes.", "answer": False}],
        "match_column": [
            {"left": "Noyau", "right": "Contrôle la cellule"},
            {"left": "Vacuole", "right": "Stocke l'eau"},
        ],
        "short_answer": [],
    }
    doc = _build_ws(payload, "fr", tmp_path,
                    book={"grade": "5e", "subject": "SVT"},
                    chap={"title": "La cellule", "sections": []})
    blob = _blob(doc)
    assert "Section A — Remplissez les blancs" in blob
    assert "Section B — Vrai ou Faux" in blob
    assert "Section C — Reliez les colonnes" in blob
    assert "Corrigé" in blob
    assert "Answer Key" not in blob
    assert "Fill in the blanks" not in blob
    secs = dict(_sections(doc))
    assert secs["Section B — Vrai ou Faux"] == ["1. Faux"]
    _match_table(doc, ["Colonne A", "Colonne B"])  # localized headers, A/B letters


# ── English: byte-identical to the pre-i18n output ───────────────────────────

def test_english_output_unchanged(tmp_path):
    payload = {
        "title": "Cells Worksheet",
        "instructions": "Answer all questions.",
        "fill_blank": [{"q": "The ____ controls the cell.", "answer": "nucleus"}],
        "true_false": [{"statement": "Animal cells have chloroplasts.", "answer": False}],
        "match_column": [
            {"left": "Nucleus", "right": "Controls the cell"},
            {"left": "Vacuole", "right": "Stores water"},
        ],
        "short_answer": [{"q": "Why are plant cells rigid?", "answer": "The cell wall.", "work_space_lines": 1}],
    }
    doc = _build_ws(payload, "en", tmp_path,
                    book={"grade": "Grade 7", "subject": "Science"},
                    chap={"title": "Cells", "sections": []})
    blob = _blob(doc)
    # The exact current literals, character for character.
    assert "Section A — Fill in the blanks" in blob
    assert "Section B — True or False" in blob
    assert "Section C — Match the columns" in blob
    assert "Section D — Short answer" in blob
    assert "Answer Key" in blob
    assert "Match each item in Column A to the correct item in Column B." in blob
    assert "Difficulty: medium" in blob  # masthead subtitle, label + raw value
    secs = dict(_sections(doc))
    assert secs["Section B — True or False"] == ["1. False"]
    _match_table(doc, ["Column A", "Column B"])


def test_letters_helper_sequences():
    assert letters("en")[:4] == ["A", "B", "C", "D"]
    assert letters("fr")[2] == "C"
    assert letters("ar")[:6] == _ABJAD_FIRST
    assert letters("ms-arab")[:6] == _ABJAD_FIRST
    assert len(letters("ar")) >= 26  # covers MAX_MATCH


def test_strings_table_complete_for_all_ten_languages():
    """Every key carries all ten languages, non-empty, with matching
    {placeholders} — _t falls back to English silently, so a dropped entry
    would ship an English string into an Arabic paper without failing."""
    import re
    from docgen.strings import _STRINGS

    langs = {"en", "ms", "ms-arab", "ar", "fr", "es", "pt", "hi", "mr", "te"}
    for key, table in _STRINGS.items():
        assert set(table) >= langs, f"{key}: missing {langs - set(table)}"
        en_ph = set(re.findall(r"{\w*}", table["en"]))
        for lang in langs:
            assert table[lang].strip(), f"{key}/{lang}: empty"
            assert set(re.findall(r"{\w*}", table[lang])) == en_ph, f"{key}/{lang}: placeholder drift"
