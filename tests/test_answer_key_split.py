"""Student/teacher document split (founder direction 2026-08-18): answer keys
stay with the teacher/parent — a student must NEVER receive a document that
contains its own answers.

exam_paper, worksheet, activity and case_study now emit TWO documents (like
docgen/exam.py always has): the student sheet first, the answer-key/teacher-
notes document second. These tests are the leak gate: each builder runs on a
fixed synthetic payload whose every answer carries a distinctive sentinel
(XQZANSWER…), then the SAVED files' XML is searched directly —

  (a) no sentinel and no localized answer-key/teacher heading appears anywhere
      in the student sheet,
  (b) every sentinel appears in the key document,
  (c) both documents build cleanly on a letterhead template stripped of Word's
      built-in styles (the 2026-08-18 production precondition — helper pattern
      from tests/test_docx_letterhead_styles.py).

Run: python -m pytest tests/test_answer_key_split.py -q
"""

from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document

from docgen import activity, case_study, exam_paper, worksheet


class _StubClient:
    """Returns a fixed payload (no network)."""

    def __init__(self, data):
        self._data = data

    def analyze(self, prompt, max_tokens=0, **k):
        return {"data": self._data}


_BOOK = {"grade": "Grade 7", "subject": "Science"}
_CHAPTER = {"title": "Cells", "sections": [{"section_title": "Cells", "content": "A cell is the basic unit."}]}


# Every answer field carries a sentinel; question/scenario text never does.
# Match rights are NOT sentinels — they print on the student sheet as options
# (the answer is the letter mapping, which lives only in the key document).
_PAYLOADS = {
    "exam_paper": {
        "title": "Cells Test",
        "instructions": "Answer everything.",
        "fill_blank": [{"q": "The ____ controls the cell.", "answer": "XQZANSWER1"}],
        "true_false": [{"statement": "Cells are alive.", "answer": True}],
        "match_column": [
            {"left": "Nucleus", "right": "Controls the cell"},
            {"left": "Vacuole", "right": "Stores water"},
        ],
        "subjective": [{"q": "Explain osmosis.", "marks": 5, "answer_outline": "XQZANSWER2"}],
    },
    "worksheet": {
        "title": "Cells Worksheet",
        "instructions": "Answer all questions.",
        "fill_blank": [{"q": "The ____ controls the cell.", "answer": "XQZANSWER1"}],
        "true_false": [{"statement": "Animal cells have chloroplasts.", "answer": False}],
        "match_column": [
            {"left": "Nucleus", "right": "Controls the cell"},
            {"left": "Vacuole", "right": "Stores water"},
        ],
        "short_answer": [{"q": "Why are plant cells rigid?", "answer": "XQZANSWER2", "work_space_lines": 1}],
    },
    "activity": {
        "intro": "XQZANSWER1 overview for the teacher.",
        "activities": [
            {"name": "Model a cell", "objective": "Understand parts", "grouping": "pairs",
             "duration_minutes": 10, "materials": ["clay"],
             "steps": ["Roll the clay", "Label the parts"],
             "teacher_facilitation": "XQZANSWER2", "success_looks_like": "XQZANSWER3"},
            {"name": "Cell walk", "objective": "Recall terms", "grouping": "whole class",
             "duration_minutes": 15, "materials": ["cards"],
             "steps": ["Hand out cards", "Discuss"],
             "teacher_facilitation": "XQZANSWER4", "success_looks_like": "XQZANSWER5"},
        ],
    },
    "case_study": {
        "title": "The Pond Case",
        "scenario": "A pond turned green overnight.",
        "background": ["The pond is shallow."],
        "discussion_questions": [
            {"q": "Why did the pond turn green?", "guidance": "XQZANSWER1"},
            {"q": "What limits the growth?", "guidance": "XQZANSWER2"},
        ],
        "concepts_applied": ["photosynthesis"],
    },
}

_PARAMS = {
    "exam_paper": {"objective": {"fill_blank": 1, "true_false": 1, "match_column": 2}, "subjective": 1},
    "worksheet": {"num_questions": 10},
    "activity": {"num_activities": 2},
    "case_study": {"num_questions": 2},
}

_BUILDERS = {
    "exam_paper": exam_paper.build,
    "worksheet": worksheet.build,
    "activity": activity.build,
    "case_study": case_study.build,
}

# Localized answer-key / teacher headings (en run): none of these may appear in
# a student sheet. "Teacher facilitation" doubles as the activity key masthead.
_FORBIDDEN_IN_STUDENT = {
    "exam_paper": ["Answer Key", "For the teacher only"],
    "worksheet": ["Answer Key", "For the teacher only"],
    "activity": ["Teacher facilitation", "Success looks like", "For the teacher only"],
    "case_study": ["Teacher notes", "For the teacher only"],
}

_SENTINELS = {
    "exam_paper": ["XQZANSWER1", "XQZANSWER2"],
    "worksheet": ["XQZANSWER1", "XQZANSWER2"],
    "activity": ["XQZANSWER1", "XQZANSWER2", "XQZANSWER3", "XQZANSWER4", "XQZANSWER5"],
    "case_study": ["XQZANSWER1", "XQZANSWER2"],
}

KINDS = sorted(_BUILDERS)


def _build(kind, out_dir, template=None):
    paths = _BUILDERS[kind](_BOOK, _CHAPTER, {}, _StubClient(_PAYLOADS[kind]),
                            _PARAMS[kind], out_dir, template=template)
    assert isinstance(paths, list) and len(paths) == 2, \
        f"{kind}: build() must return [student_sheet, key_document]"
    return paths


def _xml(path: Path) -> str:
    """The saved file's main document XML — searched raw so nothing a text
    walker might skip (table cells, split runs, field text) can hide a leak."""
    with zipfile.ZipFile(str(path)) as z:
        return z.read("word/document.xml").decode("utf-8")


@pytest.mark.parametrize("kind", KINDS)
def test_student_sheet_leaks_no_answers(kind, tmp_path):
    student, _key = _build(kind, tmp_path)
    xml = _xml(student)
    for sentinel in _SENTINELS[kind]:
        assert sentinel not in xml, f"{kind}: answer {sentinel} leaked into the student sheet"
    for heading in _FORBIDDEN_IN_STUDENT[kind]:
        assert heading not in xml, f"{kind}: {heading!r} appears in the student sheet"


@pytest.mark.parametrize("kind", KINDS)
def test_key_document_carries_every_answer(kind, tmp_path):
    _student, key = _build(kind, tmp_path)
    xml = _xml(key)
    for sentinel in _SENTINELS[kind]:
        assert sentinel in xml, f"{kind}: answer {sentinel} missing from the key document"
    # The key names itself as teacher-only material.
    assert "For the teacher only" in xml


# ── (c) both documents survive a styleless letterhead ────────────────────────
# Helper pattern from tests/test_docx_letterhead_styles.py: a valid letterhead
# whose built-in styles are absent (Word keeps them latent) — the 2026-08-18
# production failure mode. Written to a real file because build() opens the
# template once per document.

_STRIPPED = ("Title", "Heading 1", "Heading 2", "List Bullet")


def _letterhead_without_builtin_styles(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Sunrise Academy — letterhead")
    styles_el = doc.styles.element
    for name in _STRIPPED:
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        styles_el.remove(st.element)
    for name in _STRIPPED:
        with pytest.raises(KeyError):
            doc.styles[name]  # noqa: B018 — the lookup itself is the assertion
    p = tmp_path / "letterhead.docx"
    doc.save(str(p))
    return p


@pytest.mark.parametrize("kind", KINDS)
def test_split_builds_cleanly_on_styleless_letterhead(kind, tmp_path):
    template = _letterhead_without_builtin_styles(tmp_path)
    out_dir = tmp_path / "out"
    out_dir.mkdir()
    student, key = _build(kind, out_dir, template=str(template))
    # Both saved files reopen as valid documents with the fabricated styles.
    for path in (student, key):
        reopened = Document(str(path))
        assert reopened.styles["Title"] is not None
    # The split contract holds on branded documents too.
    assert all(s not in _xml(student) for s in _SENTINELS[kind])
    assert all(s in _xml(key) for s in _SENTINELS[kind])
