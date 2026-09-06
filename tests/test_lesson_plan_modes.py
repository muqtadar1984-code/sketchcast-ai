"""Lesson plan teaching modes (catalogue Phase 3, decision 11): with
``params.lesson_modes`` the prompt asks for full_lesson / micro_clip / flipped
and the document renders "Mode A — Full lesson", "Mode B — In-class
micro-clips" and "Mode C — Flipped / pre-watch" after the lesson flow, citing
clips as ``[mm:ss–mm:ss]`` computed from ``params.clips``; without the flag
the prompt and the document are what they always were.

The model is a stub that records the prompt and answers canned JSON.
"""

from __future__ import annotations

import re

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from docgen import lesson_plan
from docgen.lesson_plan import MODES_BLOCK, clip_cite, clips_block, mmss, pick_micro_clips, usable_clips


class _Stub:
    def __init__(self, data):
        self._data = data
        self.calls = []

    def analyze(self, prompt, max_tokens=0, **k):
        self.calls.append({"prompt": prompt, "max_tokens": max_tokens, **k})
        return {"data": self._data}


BOOK = {"grade": "Grade 7", "subject": "Science"}
CHAPTER = {"title": "Cells", "sections": [{"section_title": "Cells", "content": "A cell is the basic unit."}]}
CLIPS = [
    {"part": 1, "start": 0, "end": 130, "label": "What a cell is", "purpose": "introduce the cell"},
    {"part": 1, "start": 130, "end": 300.4, "label": "Structures common to all cells", "purpose": "name the organelles"},
    {"part": 2, "start": 0, "end": 180, "label": "Plant cells", "purpose": "contrast plant and animal cells"},
    {"part": 2, "start": 180, "end": 400, "label": "From cells to organisms", "purpose": "build up to organ systems"},
]
BASE = {"title": "Cells: a lesson", "duration_minutes": 45, "learning_objectives": ["State that all organisms are made of cells."],
        "materials": ["Microscope"], "key_vocabulary": [{"term": "cell", "definition": "the unit of life"}],
        "lesson_flow": [{"phase": "Hook", "minutes": 5, "teacher_does": "Shows pond water", "students_do": "Predict"}],
        "assessment": ["Exit ticket"], "homework": ["Read section 2"], "differentiation": {"support": "Word bank", "challenge": "Compare"}}
MODES = {
    "full_lesson": {"steps": ["Watch the whole video", "Run the activity in pairs", "Complete the worksheet"]},
    "micro_clip": [{"clip": 2, "task": "List the four organelles you saw."},
                   {"clip": 3, "task": "Predict which structures a root cell lacks."},
                   {"clip": 9, "task": "(no such clip)"},
                   {"clip": 3, "task": "(a repeat)"}],
    "flipped": {"pre_watch": "Watch parts 1 and 2 at home and note three structures.",
                "discussion_prompts": ["Why is the cell wall only in plants?", "What limits cell size?"],
                "activity": "Build a labelled cell model."},
}


def _iter_paras(doc):
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            for row in Table(child, doc).rows:
                for cell in row.cells:
                    yield from cell.paragraphs


def _texts(doc):
    return [re.sub(r"\s+", " ", p.text).strip() for p in _iter_paras(doc) if p.text.strip()]


def _headings(doc):
    return [re.sub(r"\s+", " ", p.text).strip() for p in _iter_paras(doc)
            if (p.style.name or "").startswith(("Heading", "Title")) and p.text.strip()]


def _build(tmp_path, params, data):
    stub = _Stub(data)
    path = lesson_plan.build(BOOK, CHAPTER, {}, stub, params, tmp_path)
    return Document(str(path)), stub


# ── pure helpers ────────────────────────────────────────────────────────


def test_mmss_and_clip_cite():
    assert mmss(0) == "00:00" and mmss(125) == "02:05" and mmss(300.4) == "05:00" and mmss(3725) == "62:05"
    assert mmss("nope") == "00:00"
    assert clip_cite(CLIPS[1]) == "[02:10–05:00]"


def test_usable_clips_drops_what_cannot_be_cited():
    clips = usable_clips({"clips": CLIPS + [{"start": 10, "end": 5}, {"start": "x", "end": 3}, "junk", {"end": 30}]})
    assert [c["label"] for c in clips] == [c["label"] for c in CLIPS]
    assert usable_clips({}) == [] and usable_clips({"clips": None}) == []


def test_clips_block_numbers_cites_and_names_parts_only_when_several():
    block = clips_block(usable_clips({"clips": CLIPS}))
    assert block.splitlines()[0] == "1. Part 1 · [00:00–02:10] What a cell is — introduce the cell"
    assert block.splitlines()[2].startswith("3. Part 2 · [00:00–03:00] Plant cells")
    one_part = clips_block(usable_clips({"clips": CLIPS[:2]}))
    assert "Part" not in one_part and one_part.startswith("1. [00:00–02:10] What a cell is")


def test_pick_micro_clips_resolves_numbers_drops_junk_and_tops_up_to_two():
    clips = usable_clips({"clips": CLIPS})
    picked = pick_micro_clips(MODES, clips)
    assert [(c["label"], t) for c, t in picked] == [
        ("Structures common to all cells", "List the four organelles you saw."),
        ("Plant cells", "Predict which structures a root cell lacks."),
    ]
    # One valid choice → topped up from the first unused clip, task = its authored purpose.
    picked = pick_micro_clips({"micro_clip": [{"clip": 4, "task": "t"}]}, clips)
    assert [(c["label"], t) for c, t in picked] == [("From cells to organisms", "t"), ("What a cell is", "introduce the cell")]
    # Never more than four.
    many = pick_micro_clips({"micro_clip": [{"clip": n, "task": "t"} for n in (1, 2, 3, 4, 1, 2)]}, clips + clips)
    assert len(many) == 4


# ── the builder ─────────────────────────────────────────────────────────


def test_with_the_flag_the_prompt_asks_for_the_modes_and_lists_the_clips(tmp_path):
    _doc, stub = _build(tmp_path, {"lesson_modes": True, "clips": CLIPS}, {**BASE, **MODES})
    (call,) = stub.calls
    prompt = call["prompt"]
    assert "TEACHING MODES" in prompt and '"full_lesson"' in prompt and '"micro_clip"' in prompt and '"flipped"' in prompt
    assert "choose 2-4 of the clips" in prompt
    assert "1. Part 1 · [00:00–02:10] What a cell is — introduce the cell" in prompt
    assert "4. Part 2 · [03:00–06:40] From cells to organisms — build up to organ systems" in prompt
    assert call["max_tokens"] > 3000, "three more sections need room"


def test_with_the_flag_three_mode_headings_render_after_the_lesson_flow_citing_clips(tmp_path):
    doc, _stub = _build(tmp_path, {"lesson_modes": True, "clips": CLIPS, "curriculum_header": ["Cambridge 0893 · 7Bs.01"]},
                        {**BASE, **MODES})
    heads = _headings(doc)
    modes = [h for h in heads if h.startswith("Mode ")]
    assert modes == ["Mode A — Full lesson", "Mode B — In-class micro-clips", "Mode C — Flipped / pre-watch"]
    assert heads.index("Lesson flow") < heads.index("Mode A — Full lesson") < heads.index("Assessment")
    texts = _texts(doc)
    blob = "\n".join(texts)
    # Mode A: the steps, numbered.
    assert "1. Watch the whole video" in blob and "3. Complete the worksheet" in blob
    # Mode B: the chosen clips cited from THEIR start/end, then the task; junk numbers gone.
    assert "1. Part 1 · [02:10–05:00] Structures common to all cells" in blob
    assert "2. Part 2 · [00:00–03:00] Plant cells" in blob
    assert "List the four organelles you saw." in blob and "(no such clip)" not in blob and "(a repeat)" not in blob
    assert len(re.findall(r"\[\d\d:\d\d–\d\d:\d\d\]", blob)) == 2
    # Mode C: pre-watch, prompts, activity.
    assert "Before class: Watch parts 1 and 2 at home and note three structures." in blob
    assert "Why is the cell wall only in plants?" in blob and "Activity: Build a labelled cell model." in blob
    # The curriculum block rides along (decision 10).
    assert "Cambridge 0893 · 7Bs.01" in blob


def test_without_the_flag_the_prompt_and_document_are_unchanged(tmp_path):
    doc, stub = _build(tmp_path, {"duration_minutes": 45}, {**BASE, **MODES})
    (call,) = stub.calls
    prompt = call["prompt"]
    assert "TEACHING MODES" not in prompt and "full_lesson" not in prompt and "micro_clip" not in prompt
    assert call["max_tokens"] == 3000
    assert prompt.rstrip().endswith("4-6 lesson_flow phases that sum to roughly 45 minutes.")
    heads = _headings(doc)
    assert not [h for h in heads if h.startswith("Mode ")]
    # Even a reply that VOLUNTEERS the mode fields renders nothing for them.
    assert "[02:10–05:00]" not in "\n".join(_texts(doc))
    assert heads == ["Cells: a lesson", "Learning objectives", "Materials", "Key vocabulary", "Lesson flow",
                     "Assessment", "Homework", "Differentiation"]


def test_the_flag_without_any_usable_clip_renders_the_plain_plan(tmp_path):
    doc, stub = _build(tmp_path, {"lesson_modes": True, "clips": [{"start": 5, "end": 5}]}, {**BASE, **MODES})
    assert "TEACHING MODES" not in stub.calls[0]["prompt"]
    assert not [h for h in _headings(doc) if h.startswith("Mode ")]


def test_localized_mode_headings_use_the_language_s_letters(tmp_path):
    doc_ar = Document(str(lesson_plan.build(BOOK, CHAPTER, {}, _Stub({**BASE, **MODES}),
                                            {"lesson_modes": True, "clips": CLIPS}, tmp_path, language="ar")))
    ar_modes = [h for h in _headings(doc_ar) if h.startswith("النمط ")]
    assert ar_modes == ["النمط أ — الدرس الكامل", "النمط ب — مقاطع قصيرة داخل الصف", "النمط ج — الصف المعكوس / مشاهدة مسبقة"]
    assert MODES_BLOCK.count("{clips}") == 1
