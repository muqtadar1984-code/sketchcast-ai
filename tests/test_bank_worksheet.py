"""``docgen.bank_worksheet.write_set``: approved bank items → a student sheet
and a SEPARATE answer key. The leak gate is the same as
tests/test_answer_key_split.py's — every answer carries a sentinel and the
saved files' XML is searched directly — plus the layout rules of decision 9:
curriculum block, one section per type in TYPE_ORDER on its own page, MCQ
options A–D, the match table shuffled by the set's seed, ruled lines sized
by est_seconds/marks, the total-marks line.
"""

from __future__ import annotations

import random
import re
import zipfile

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from catalogue import questions
from docgen import bank_worksheet as bw
from docgen.strings import letters

HEADER = ["Cambridge Lower Secondary Science 0893 · 7Bs.01, 7Bs.02", "CBSE Science 086 · Class 9 · Cell — the basic unit of life"]

# Every answer carries an XQZ sentinel; stems and options never do.
ITEMS = [
    {"id": "q-mcq", "item_type": "mcq", "difficulty": 2, "marks": 1, "est_seconds": 60,
     "stem": "Which structure controls what enters and leaves the cell?",
     "options": [{"key": "A", "text": "Nucleus"}, {"key": "B", "text": "Cell membrane"}, {"key": "C", "text": "Cytoplasm"}, {"key": "D", "text": "Mitochondrion"}],
     "answer": {"key": "B"}, "marking_scheme": [{"point": "B", "marks": 1}], "explanation": "XQZEXPL1 the membrane is the boundary."},
    {"id": "q-tf", "item_type": "true_false", "difficulty": 1, "marks": 1, "est_seconds": 30,
     "stem": "Animal cells have a cell wall.", "answer": {"value": False}, "marking_scheme": [], "explanation": None},
    {"id": "q-fill", "item_type": "fill_blank", "difficulty": 1, "marks": 1, "est_seconds": 45,
     "stem": "The ____ holds the genetic material.", "answer": {"text": "XQZNUCLEUS", "accept": ["XQZALT"]}, "marking_scheme": []},
    {"id": "q-match", "item_type": "match", "difficulty": 2, "marks": 3, "est_seconds": 90,
     "stem": "Match each structure to its job.",
     "options": {"pairs": [{"left": "Nucleus", "right": "Controls the cell"}, {"left": "Chloroplast", "right": "Photosynthesis"}, {"left": "Vacuole", "right": "Stores water"}]},
     "answer": {"pairs": [{"left": "Nucleus", "right": "Controls the cell"}, {"left": "Chloroplast", "right": "Photosynthesis"}, {"left": "Vacuole", "right": "Stores water"}]},
     "marking_scheme": []},
    {"id": "q-ar", "item_type": "assertion_reason", "difficulty": 4, "marks": 1, "est_seconds": 75,
     "stem": "Assertion (A): Root cells have no chloroplasts. Reason (R): Roots never see light.",
     "options": [{"key": "A", "text": "Both true; R explains A"}, {"key": "B", "text": "Both true; R does not explain A"}, {"key": "C", "text": "A true, R false"}, {"key": "D", "text": "A false, R true"}],
     "answer": {"key": "A"}, "marking_scheme": []},
    {"id": "q-short", "item_type": "short_answer", "difficulty": 3, "marks": 2, "est_seconds": 120,
     "stem": "Explain why a muscle cell has many mitochondria.",
     "answer": {"text": "XQZSHORT it respires a lot."}, "marking_scheme": [{"point": "XQZPOINT1 respiration", "marks": 1}, {"point": "XQZPOINT2 energy demand", "marks": 1}],
     "explanation": "XQZEXPL2"},
    {"id": "q-long", "item_type": "long_answer", "difficulty": 4, "marks": 5, "est_seconds": 480,
     "stem": "Compare plant and animal cells.", "answer": {"text": "XQZLONG"}, "marking_scheme": [{"point": "XQZWALL", "marks": 2}, {"point": "XQZCHLORO", "marks": 3}]},
    {"id": "q-num", "item_type": "numerical", "difficulty": 3, "marks": 3, "est_seconds": 180,
     "stem": "A cell is 0.02 mm wide. How many fit across 1 mm?", "answer": {"value": 50, "unit": "cells", "tolerance": 0},
     "marking_scheme": [{"point": "XQZMETHOD divide", "marks": 2}, {"point": "answer", "marks": 1}]},
    {"id": "q-diag", "item_type": "diagram_label", "difficulty": 2, "marks": 3, "est_seconds": 150,
     "stem": "The plant cell figure: part 1 is the rigid outer layer, part 2 is green, part 3 is the large central space.",
     "options": {"figure_key": "plant_cell", "caption": "A plant cell"},
     "answer": {"labels": [{"n": 1, "label": "XQZWALLLBL"}, {"n": 2, "label": "XQZCHLOROLBL"}, {"n": 3, "label": "XQZVACLBL"}]},
     "marking_scheme": []},
]
SENTINELS = sorted(set(re.findall(r"XQZ\w+", repr(ITEMS))))


def _xml(path):
    with zipfile.ZipFile(path) as z:
        return z.read("word/document.xml").decode("utf-8")


def _iter_paras(doc):
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            for row in Table(child, doc).rows:
                for cell in row.cells:
                    yield from cell.paragraphs


def _texts(doc):
    return [re.sub(r"\s+", " ", p.text).strip() for p in _iter_paras(doc) if p.text.strip()]


def _headings(doc):
    return [re.sub(r"\s+", " ", p.text).strip() for p in _iter_paras(doc)
            if (p.style.name or "").startswith("Heading") and p.text.strip()]


def _page_breaks(path):
    return len(re.findall(r'<w:br [^>]*w:type="page"', _xml(path)))


def _write(tmp_path, items=ITEMS, seed=7, language="en", **kw):
    return bw.write_set(tmp_path, "Cells · Quick check", "9 questions · Total marks: 20", items,
                        header_lines=HEADER, language=language, template=None, blueprint_name="Quick check",
                        rng=random.Random(seed), **kw)


def test_type_order_is_the_bank_s_item_type_enum():
    assert bw.TYPE_ORDER == questions.ITEM_TYPES
    assert bw.SUBJECTIVE_TYPES == questions.SUBJECTIVE_TYPES


def test_two_files_header_lines_and_no_answer_in_the_student_sheet(tmp_path):
    sheet, key = _write(tmp_path)
    assert sheet.exists() and key.exists() and sheet != key
    sheet_xml, key_xml = _xml(sheet), _xml(key)
    for line in HEADER:
        assert line in sheet_xml and line in key_xml
    for s in SENTINELS:
        assert s not in sheet_xml, s
        assert s in key_xml, s
    assert "Answer Key" not in sheet_xml and "Answer Key" in key_xml
    assert "For the teacher only" in key_xml and "For the teacher only" not in sheet_xml
    # True/False and MCQ letter answers are not sentinels; they must still be absent from the sheet.
    sheet_heads = _headings(Document(str(sheet)))
    assert sheet_heads == [f"Section {L} — {name}" for L, name in zip("ABCDEFGHI", (
        "Multiple choice", "True or False", "Fill in the blanks", "Match the columns", "Assertion and reason",
        "Short answer", "Long answer", "Numerical", "Label the diagram"))]


def test_page_breaks_between_sections_and_none_elsewhere(tmp_path):
    sheet, key = _write(tmp_path)
    sections = len(bw.group_items(ITEMS))
    assert sections == 9
    assert _page_breaks(sheet) == sections - 1
    # A subset: three types present → two breaks; one type → none.
    (tmp_path / "three").mkdir()
    sheet3, _ = _write(tmp_path / "three", items=[ITEMS[0], ITEMS[5], ITEMS[7]])
    assert _page_breaks(sheet3) == 2
    (tmp_path / "one").mkdir()
    sheet1, _ = _write(tmp_path / "one", items=[ITEMS[5], ITEMS[5]])
    assert _page_breaks(sheet1) == 0


def test_mcq_options_are_lettered_and_the_key_gives_only_the_letter(tmp_path):
    sheet, key = _write(tmp_path)
    texts = _texts(Document(str(sheet)))
    i = texts.index("1. Which structure controls what enters and leaves the cell? [1]")
    assert texts[i + 1:i + 5] == ["A) Nucleus", "B) Cell membrane", "C) Cytoplasm", "D) Mitochondrion"]
    key_texts = _texts(Document(str(key)))
    k = key_texts.index("Section A — Multiple choice")
    assert key_texts[k + 1] == "1. B"
    assert "Cell membrane" not in "\n".join(key_texts), "option text is student-facing; the key gives the letter"
    assert "Explanation: XQZEXPL1 the membrane is the boundary." in key_texts


def test_match_table_is_shuffled_by_the_seed_and_the_key_maps_the_letters(tmp_path):
    sheet, key = _write(tmp_path, seed=3)
    doc = Document(str(sheet))
    table = next(t for t in doc.tables if [c.text for c in t.rows[0].cells] == ["Column A", "Column B"])
    rows = [[c.text for c in r.cells] for r in table.rows[1:]]
    assert [r[0] for r in rows] == ["1. Nucleus", "2. Chloroplast", "3. Vacuole"]
    right_by_letter = {r[1].split(". ", 1)[0]: r[1].split(". ", 1)[1] for r in rows}
    key_line = next(t for t in _texts(Document(str(key))) if t.startswith("1. 1 → "))
    mapping = dict(re.findall(r"(\d) → ([A-Z])", key_line))
    pairs = ITEMS[3]["options"]["pairs"]
    for n, letter in mapping.items():
        assert right_by_letter[letter] == pairs[int(n) - 1]["right"]
    # Deterministic in the seed: same seed → same table; the seed is what a re-render keys on.
    (tmp_path / "again").mkdir()
    sheet2, _ = _write(tmp_path / "again", seed=3)
    table2 = next(t for t in Document(str(sheet2)).tables if [c.text for c in t.rows[0].cells] == ["Column A", "Column B"])
    assert [[c.text for c in r.cells] for r in table2.rows] == [[c.text for c in r.cells] for r in table.rows]
    orders = set()
    for seed in range(6):
        d = tmp_path / f"s{seed}"
        d.mkdir()
        s, _ = _write(d, seed=seed)
        t = next(t for t in Document(str(s)).tables if [c.text for c in t.rows[0].cells] == ["Column A", "Column B"])
        orders.add(tuple(r.cells[1].text for r in t.rows[1:]))
    assert len(orders) > 1, "different seeds shuffle differently"


def test_subjective_items_get_ruled_lines_sized_by_time_and_marks(tmp_path):
    assert bw.answer_lines({"marks": 2, "est_seconds": 120}) == 4       # 2 lines a mark beats 3 by time
    assert bw.answer_lines({"marks": 5, "est_seconds": 480}) == 12      # time wins
    assert bw.answer_lines({"marks": 1, "est_seconds": 10}) == 2        # floor
    assert bw.answer_lines({"marks": 20, "est_seconds": 5000}) == 16    # ceiling
    assert bw.answer_lines({}) == 4
    sheet, _ = _write(tmp_path)
    xml = _xml(sheet)
    # Ruled lines are bordered empty paragraphs; the short (4) + long (12) + numerical (6) items = 22.
    assert len(re.findall(r'<w:bottom w:val="single" w:sz="6"', xml)) == 4 + 12 + 6


def test_total_marks_line_and_diagram_blanks(tmp_path):
    sheet, key = _write(tmp_path)
    texts = _texts(Document(str(sheet)))
    assert "Quick check · 9 questions · Total marks: 20" in texts
    assert bw.total_marks(ITEMS) == 20
    i = texts.index("1. The plant cell figure: part 1 is the rigid outer layer, part 2 is green, part 3 is the large central space. [3]")
    assert texts[i + 1:i + 4] == ["1. ________________", "2. ________________", "3. ________________"]
    key_texts = _texts(Document(str(key)))
    assert "1. 1 XQZWALLLBL; 2 XQZCHLOROLBL; 3 XQZVACLBL" in key_texts
    assert "1. 50 cells" in key_texts
    assert "Marking scheme: XQZMETHOD divide (2); answer (1)" in key_texts
    assert "1. XQZNUCLEUS / XQZALT" in key_texts
    assert "1. False" in key_texts


def test_arabic_sheet_uses_abjad_letters_for_options_and_sections(tmp_path):
    sheet, key = _write(tmp_path, language="ar")
    texts = _texts(Document(str(sheet)))
    ab = letters("ar")
    assert f"{ab[0]}) Nucleus" in texts and f"{ab[3]}) Mitochondrion" in texts
    assert any(t.startswith(f"القسم {ab[0]} — ") for t in texts)
    assert "Answer Key" not in _xml(sheet)


def test_unknown_item_types_are_skipped_not_rendered(tmp_path):
    sheet, _ = _write(tmp_path, items=ITEMS[:1] + [{"item_type": "essay", "stem": "XQZODD"}])
    assert "XQZODD" not in _xml(sheet)
    assert bw.group_items([{"item_type": "essay"}]) == []


def test_no_rng_means_the_fixed_seed_zero_shuffle(tmp_path):
    def column_b(sub, rng):
        (tmp_path / sub).mkdir()
        paths = bw.write_set(tmp_path / sub, "T", "S", [ITEMS[3]], header_lines=None, language="en", template=None,
                             blueprint_name="", rng=rng)
        table = next(t for t in Document(str(paths[0])).tables if [c.text for c in t.rows[0].cells] == ["Column A", "Column B"])
        return [r.cells[1].text for r in table.rows[1:]]

    assert column_b("none", None) == column_b("zero", random.Random(0))
    assert sorted(t.split(". ", 1)[1] for t in column_b("again", None)) == ["Controls the cell", "Photosynthesis", "Stores water"]
