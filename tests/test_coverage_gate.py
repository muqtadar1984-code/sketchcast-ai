"""Coverage gate — the Stage 0 check that must exist BEFORE any model downgrade.

The analyzer publishes, per episode, the sections it chunked and the concepts it
extracted; the script prompt is built from exactly those two lists; and until now
nothing compared the reply back to them. A script that dropped half a chapter
finished as a clean success. That is survivable on Sonnet and dangerous the
moment a cheaper model is introduced, because thinning is how a cheaper model
fails and it fails silently.

Thresholds under test were calibrated on the one real Sonnet analysis+script pair
(23 topics, 14 segments): the script as generated measured 0.913, a third of its
topics never mentioned measured 0.652, two thirds of its segments deleted
measured 0.348. See shared/coverage.py for the full table.
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest

from shared.coverage import (chapter_topics, docx_text, gate_mode, measure,
                             script_text, should_fail)

PROJECT_ROOT = Path(__file__).resolve().parent.parent


def _analysis(names, sections=(), episodes=1):
    """A MasterAnalysis-shaped dict whose episode carries every concept."""
    concepts = [
        {"concept_id": f"c{i + 1:03d}", "name": n, "definition": f"{n} explained."}
        for i, n in enumerate(names)
    ]
    return {
        "chapter_title": "Chapter",
        "concepts": {"concepts": concepts, "dependencies": [], "prerequisites": []},
        "episodes": {
            "chapter_title": "Chapter",
            "total_episodes": episodes,
            "episodes": [{
                "episode_num": 1,
                "title": "Chapter",
                "sections_covered": list(sections),
                "key_concepts_introduced": [c["concept_id"] for c in concepts],
            }],
        },
    }


def _script(*paragraphs):
    """A ChapterScripts-shaped episode: one narration segment per paragraph."""
    return {"segments": [{"type": "explore", "text": p} for p in paragraphs]}


# The twelve topics used by the covers-everything / drops-a-third pair below.
# Deliberately a real chapter's shape: distinct technical terms, plus a sibling
# family that shares a head word ("cell wall" / "cell membrane").
TOPICS = [
    "Photosynthesis", "Chlorophyll", "Stomata", "Cell wall", "Cell membrane",
    "Chloroplast", "Glucose", "Respiration", "Xylem", "Phloem",
    "Transpiration", "Osmosis",
]

FULL_SCRIPT = _script(
    "Every green plant runs photosynthesis, and the pigment that makes it possible "
    "is chlorophyll, packed inside each chloroplast.",
    "Air enters through the stomata on the underside of the leaf. Water arrives by "
    "osmosis and is pulled up the xylem, a journey driven by transpiration.",
    "The sugar the leaf builds is glucose, and the phloem carries it away to the rest "
    "of the plant.",
    "A plant cell is held rigid by its cell wall, while the cell membrane underneath "
    "decides what actually gets in and out.",
    "At night the same glucose is broken down again in respiration, which is "
    "photosynthesis running backwards.",
)

# The same lesson with a third of the topics simply never mentioned — no error,
# no truncation, full-length fluent prose. This is the failure mode.
THIN_SCRIPT = _script(
    "Every green plant runs photosynthesis, and the pigment that makes it possible "
    "is chlorophyll, packed inside each chloroplast.",
    "Air enters through the stomata on the underside of the leaf.",
    "The sugar the leaf builds is glucose, and it powers everything the plant does. "
    "Later the plant breaks that sugar down again in respiration.",
    "A plant cell is held rigid by its cell wall, and the cell membrane just inside "
    "it decides what gets through. Cells are the building blocks of every leaf, "
    "every root and every stem, and they work together beautifully.",
)


class TestCoversEverything:
    def test_a_script_that_teaches_every_topic_passes(self):
        report = measure(_analysis(TOPICS), _analysis(TOPICS)["episodes"]["episodes"][0],
                         script_text(FULL_SCRIPT))
        assert report["checked"] is True
        assert report["topics"] == 12
        assert report["missed"] == []
        assert report["covered"] == 1.0
        assert report["verdict"] == "ok"
        assert should_fail(report, "warn") is False
        assert should_fail(report, "strict") is False

    def test_inflection_and_paraphrase_are_not_misses(self):
        # "photosynthetic", "chloroplasts", "stomatal" — a script that teaches
        # the topic in the teacher's words, not the analyzer's.
        analysis = _analysis(["Photosynthesis", "Chloroplast", "Stomata"])
        script = _script(
            "Photosynthetic tissue is stuffed with chloroplasts, and stomatal pores "
            "let the air in."
        )
        report = measure(analysis, analysis["episodes"]["episodes"][0], script_text(script))
        assert report["covered"] == 1.0

    def test_a_topic_taught_only_on_the_slide_still_counts(self):
        # A concept introduced as a slide heading or a diagram node has been
        # taught; scoring narration alone would mark it missed.
        analysis = _analysis(["Osmosis"])
        script = {"segments": [{
            "type": "explore", "text": "Water moves. Let's see how.",
            "slide_heading": "Osmosis",
            "slide_points": ["water crosses the membrane"],
            "slide_visual": {"kind": "flow", "nodes": ["soil", "root"], "caption": "osmosis"},
        }]}
        report = measure(analysis, analysis["episodes"]["episodes"][0], script_text(script))
        assert report["covered"] == 1.0


class TestDropsAThirdOfTheSections:
    def test_a_third_dropped_is_measured_as_a_shortfall(self):
        analysis = _analysis(TOPICS)
        report = measure(analysis, analysis["episodes"]["episodes"][0], script_text(THIN_SCRIPT))
        # Xylem, phloem, transpiration and osmosis are never mentioned.
        assert report["topics"] == 12
        assert report["addressed"] == 8
        assert report["covered"] == pytest.approx(0.667, abs=0.001)
        assert set(report["missed"]) == {"Xylem", "Phloem", "Transpiration", "Osmosis"}
        assert report["verdict"] == "short"

    def test_a_third_dropped_does_not_burn_the_teachers_credit_by_default(self):
        # The default (warn) records the shortfall and lets the lesson through:
        # production has no coverage distribution at all yet, and hard-failing
        # on an unmeasured number costs a teacher a credit for a lesson that is
        # thin, not broken. strict — for after the flip, once baselines exist —
        # fails the same report.
        analysis = _analysis(TOPICS)
        report = measure(analysis, analysis["episodes"]["episodes"][0], script_text(THIN_SCRIPT))
        assert should_fail(report, "warn") is False
        assert should_fail(report, "strict") is True

    def test_a_collapsed_script_hits_the_floor_in_every_mode(self):
        # Below 0.30 seven topics in ten are never mentioned — not thinning,
        # a script that is not about this chapter. Same class as the
        # zero-video-parts case the worker already raises on.
        analysis = _analysis(TOPICS)
        script = _script("Plants are interesting. Photosynthesis is a thing that happens.")
        report = measure(analysis, analysis["episodes"]["episodes"][0], script_text(script))
        assert report["covered"] < 0.30
        assert report["verdict"] == "floor"
        assert should_fail(report, "warn") is True
        assert should_fail(report, "strict") is True
        assert should_fail(report, "off") is False


class TestWhatCountsAsATopic:
    def test_the_content_placeholder_is_never_a_topic(self):
        # 143 of 304 production parts carry the structurer's "Content"
        # placeholder as their only heading. Measuring a script against the word
        # "Content" is not a coverage test.
        analysis = _analysis(["Osmosis"], sections=["Content", "1.", "Water movement"])
        names = [t["name"] for t in chapter_topics(analysis, analysis["episodes"]["episodes"][0])]
        assert "Content" not in names
        assert "1." not in names
        assert "Water movement" in names

    def test_real_section_headings_are_measured_alongside_concepts(self):
        analysis = _analysis(["Osmosis"], sections=["Water movement in roots"])
        report = measure(analysis, analysis["episodes"]["episodes"][0],
                         script_text(_script("Osmosis moves water. Roots take it up.")))
        assert report["sections"] == 1 and report["concepts"] == 1
        # The heading's own words ("water movement roots") are all present.
        assert report["covered"] == 1.0

    def test_concepts_are_scoped_to_the_episode(self):
        # Part 2 must not be marked down for failing to teach Part 1's concepts.
        analysis = _analysis(["Xylem", "Phloem"])
        episode = dict(analysis["episodes"]["episodes"][0], key_concepts_introduced=["c002"])
        names = [t["name"] for t in chapter_topics(analysis, episode)]
        assert names == ["Phloem"]

    def test_a_legacy_episode_with_no_concept_ids_still_measures(self):
        # segmenter.build_single_episode and every pre-chunking cached analysis
        # leave key_concepts_introduced empty. Scoping to an empty set would
        # switch the gate off for exactly those rows.
        analysis = _analysis(["Xylem", "Phloem"])
        episode = dict(analysis["episodes"]["episodes"][0], key_concepts_introduced=[])
        names = {t["name"] for t in chapter_topics(analysis, episode)}
        assert names == {"Xylem", "Phloem"}

    def test_a_document_is_measured_against_the_whole_chapter(self):
        # episode=None — documents are generated per CHAPTER, never per part.
        analysis = _analysis(["Xylem", "Phloem"], sections=["Transport"])
        names = {t["name"] for t in chapter_topics(analysis, None)}
        assert names == {"Xylem", "Phloem", "Transport"}


class TestTheSharedHeadWordTrap:
    def test_one_mention_of_the_head_word_does_not_cover_its_siblings(self):
        # "cell membrane" / "cell wall" / "cell division" all contain "cell".
        # A rule that scores a hit on any single token marks a script that only
        # ever says "cell" as covering all three.
        analysis = _analysis(["Cell membrane", "Cell wall", "Cell division"])
        report = measure(analysis, analysis["episodes"]["episodes"][0],
                         script_text(_script("A cell is the unit of life. Every cell matters.")))
        assert report["addressed"] == 0
        assert report["covered"] == 0.0

    def test_the_sibling_actually_taught_is_the_one_credited(self):
        analysis = _analysis(["Cell membrane", "Cell wall", "Cell division"])
        report = measure(analysis, analysis["episodes"]["episodes"][0],
                         script_text(_script("The rigid cell wall is what holds a plant up.")))
        assert report["addressed"] == 1
        assert set(report["missed"]) == {"Cell membrane", "Cell division"}


class TestNonLatinScripts:
    def test_hindi_is_tokenised_not_shattered(self):
        # A \\w+ / [^\\W\\d_]+ tokeniser splits Devanagari at every vowel sign
        # (categories Mc/Mn are not alnum), shattering these two words into 15
        # fragments and scoring every Hindi lesson we ship at near zero.
        analysis = _analysis(["प्रकाश संश्लेषण", "श्वसन"])
        script = _script("पौधों में प्रकाश संश्लेषण होता है और रात में श्वसन होता है।")
        report = measure(analysis, analysis["episodes"]["episodes"][0], script_text(script))
        assert report["covered"] == 1.0

    def test_telugu_missing_topics_are_still_detected(self):
        analysis = _analysis(["కిరణజన్య సంయోగక్రియ", "శ్వాసక్రియ"])
        script = _script("మొక్కలలో కిరణజన్య సంయోగక్రియ జరుగుతుంది.")
        report = measure(analysis, analysis["episodes"]["episodes"][0], script_text(script))
        assert report["addressed"] == 1
        assert report["missed"] == ["శ్వాసక్రియ"]

    def test_arabic_is_measured(self):
        analysis = _analysis(["التمثيل الضوئي", "التنفس"])
        script = _script("تقوم النباتات بعملية التمثيل الضوئي في أوراقها.")
        report = measure(analysis, analysis["episodes"]["episodes"][0], script_text(script))
        assert report["addressed"] == 1
        assert report["missed"] == ["التنفس"]


class TestUnjudgeable:
    def test_an_analysis_with_nothing_to_measure_is_not_a_failure(self):
        # No concepts extracted and the only heading is a placeholder. Nothing
        # to compare against is not the same as bad — the same principle
        # book_health.text_quality applies to a book with too little text.
        analysis = _analysis([], sections=["Content"])
        report = measure(analysis, analysis["episodes"]["episodes"][0], script_text(_script("Hello.")))
        assert report["checked"] is False
        assert report["covered"] is None
        assert report["verdict"] == "unmeasured"
        assert should_fail(report, "strict") is False

    def test_too_few_topics_to_gate_is_still_recorded(self):
        # With 3 topics one honest miss reads as 0.67 and two as 0.33 — a good
        # short lesson would trip the floor on rounding, so it is measured and
        # recorded but never allowed to fail a job.
        analysis = _analysis(["Xylem", "Phloem", "Stomata"])
        report = measure(analysis, analysis["episodes"]["episodes"][0],
                         script_text(_script("Nothing relevant here at all.")))
        assert report["covered"] == 0.0
        assert report["verdict"] == "floor"
        assert report["gated"] is False
        assert should_fail(report, "strict") is False


class TestGateMode:
    def test_default_is_warn(self, monkeypatch):
        monkeypatch.delenv("COVERAGE_GATE", raising=False)
        assert gate_mode() == "warn"

    def test_junk_falls_back_to_warn(self, monkeypatch):
        monkeypatch.setenv("COVERAGE_GATE", "yes-please")
        assert gate_mode() == "warn"

    def test_modes_are_read_from_the_env(self, monkeypatch):
        for value in ("off", "warn", "strict"):
            monkeypatch.setenv("COVERAGE_GATE", value.upper())
            assert gate_mode() == value


class TestDocumentText:
    def test_table_cells_are_read(self, tmp_path):
        # A worksheet's matching exercise and an exam's mark scheme live
        # entirely inside tables — a paragraph-only read scores them empty.
        import docx

        doc = docx.Document()
        doc.add_paragraph("Worksheet — Transport in plants")
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Xylem"
        table.rows[0].cells[1].text = "carries water"
        path = tmp_path / "w.docx"
        doc.save(str(path))

        text = docx_text(path)
        assert "Xylem" in text and "carries water" in text


class TestRetryChannel:
    def test_missed_topics_reach_the_retry_prompt(self):
        # A retry with the identical prompt is a coin flip; the gate's retry
        # names what was dropped, so it has to actually arrive in the context.
        from agent3_scripts.script_generator import generate_episode_script

        seen: dict = {}

        class _Stub:
            model = "claude-haiku-4-5"

            def analyze(self, prompt, system=None, max_tokens=0, **kw):
                seen["prompt"] = prompt
                return {"data": {"segments": [{"type": "explore", "text": "Xylem carries water."}]}}

        analysis = _analysis(["Xylem", "Phloem"])
        generate_episode_script(
            analysis["episodes"]["episodes"][0], analysis, 1, _Stub(),
            must_cover=["Phloem", "Transpiration"],
        )
        assert "MANDATORY COVERAGE" in seen["prompt"]
        assert "Phloem" in seen["prompt"] and "Transpiration" in seen["prompt"]

    def test_no_must_cover_leaves_the_prompt_untouched(self):
        from agent3_scripts.script_generator import generate_episode_script

        seen: dict = {}

        class _Stub:
            model = "claude-sonnet-4-6"

            def analyze(self, prompt, system=None, max_tokens=0, **kw):
                seen["prompt"] = prompt
                return {"data": {"segments": [{"type": "explore", "text": "Xylem carries water."}]}}

        analysis = _analysis(["Xylem"])
        generate_episode_script(analysis["episodes"]["episodes"][0], analysis, 1, _Stub())
        assert "MANDATORY COVERAGE" not in seen["prompt"]


# ── Calibration lock ─────────────────────────────────────────────────────────
# storage/ is gitignored, so this pair only exists on a machine that has run the
# pipeline. Where it does exist it pins the numbers quoted in shared/coverage.py
# to the real Sonnet output they were measured from — change the tokeniser or a
# threshold and this is what notices.
_REAL_ANALYSIS = PROJECT_ROOT / "storage" / "analysis" / "t" / "chapter_0.json"
_REAL_SCRIPT = PROJECT_ROOT / "storage" / "scripts" / "t" / "chapter_0_episode_1.json"


@pytest.mark.skipif(
    not (_REAL_ANALYSIS.exists() and _REAL_SCRIPT.exists()),
    reason="no locally generated analysis+script pair (storage/ is gitignored)",
)
def test_the_real_sonnet_script_measures_where_the_thresholds_assume():
    analysis = json.loads(_REAL_ANALYSIS.read_text(encoding="utf-8"))
    script = json.loads(_REAL_SCRIPT.read_text(encoding="utf-8"))
    episode = analysis["episodes"]["episodes"][0]

    whole = measure(analysis, episode, script_text(script))
    assert whole["topics"] == 23
    assert whole["covered"] == pytest.approx(0.913, abs=0.001)
    assert whole["verdict"] == "ok"

    segments = script["segments"]
    n = len(segments)
    # Two thirds of the segments deleted — the worst degradation reproducible
    # from real output. It must land above the floor's 0.30, or the floor is
    # firing on something a merely-thin script can do.
    third = measure(analysis, episode, script_text({"segments": segments[: n // 3]}))
    assert third["covered"] == pytest.approx(0.348, abs=0.001)
    assert third["covered"] > 0.30
