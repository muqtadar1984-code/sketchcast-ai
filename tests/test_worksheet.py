"""Worksheet generator — questions must be TYPED and GROUPED BY KIND, and the
match-the-columns exercise must render as a TABLE (not running text). Regression
for the teacher feedback (2026-07-10): "1 fill-in-the-blank followed by 1
subjective ... match the column populated as running text."
"""

from __future__ import annotations

import json
import re

from docx import Document

from docgen import questions, worksheet


class _StubClient:
    """Returns a fixed typed worksheet payload (no network)."""

    def __init__(self, data):
        self._data = data

    def analyze(self, prompt, max_tokens=0, **k):
        return {"data": self._data}


_PAYLOAD = {
    "title": "Cells Worksheet",
    "instructions": "Answer all questions.",
    "fill_blank": [
        {"q": "The ____ controls the cell.", "answer": "nucleus"},
        {"q": "Plant cells have a ____ wall.", "answer": "cell"},
    ],
    "true_false": [{"statement": "Animal cells have chloroplasts.", "answer": False}],
    "match_column": [
        {"left": "Nucleus", "right": "Controls the cell"},
        {"left": "Chloroplast", "right": "Captures sunlight"},
        {"left": "Vacuole", "right": "Stores water"},
    ],
    "short_answer": [{"q": "Why are plant cells rigid?", "answer": "The cell wall.", "work_space_lines": 2}],
}


def test_worksheet_docx_groups_by_type_and_tables_the_match(tmp_path):
    out = worksheet.build(
        book={"grade": "Grade 7", "subject": "Science"},
        chapter={"title": "Cells", "sections": [{"section_title": "Cells", "content": "A cell is the basic unit."}]},
        analysis={},
        client=_StubClient(_PAYLOAD),
        params={"num_questions": 10, "include_answer_key": True},
        out_dir=tmp_path,
    )
    doc = Document(str(out))
    headings = [p.text for p in doc.paragraphs if p.text.strip()]
    blob = "\n".join(headings)
    # One grouped section per kind, in order.
    assert "Fill in the blanks" in blob
    assert "True or False" in blob
    assert "Match the columns" in blob
    assert "Short answer" in blob
    # The match exercise is a real 2-column TABLE, not running text.
    assert len(doc.tables) == 1
    t = doc.tables[0]
    assert [c.text for c in t.rows[0].cells] == ["Column A", "Column B"]
    assert len(t.rows) == 1 + 3  # header + 3 pairs

    # Structured questions.json is typed + grouped (fill → tf → match → short).
    payload = json.loads((tmp_path / "questions.json").read_text(encoding="utf-8"))
    types = [q["type"] for q in payload["questions"]]
    assert types == ["fill_blank", "fill_blank", "true_false", "match", "short"]
    match_q = next(q for q in payload["questions"] if q["type"] == "match")
    assert len(match_q["pairs"]) == 3


# Trips every numbering trap at once: an LLM-self-numbered fill question, a
# blank answer (must keep its slot in the key), and a "Q1)"-prefixed short answer.
_NUMBERING_PAYLOAD = {
    "title": "Numbering Worksheet",
    "instructions": "Answer all questions.",
    "fill_blank": [
        {"q": "1. The ____ controls the cell.", "answer": "nucleus"},
        {"q": "Plant cells have a ____ wall.", "answer": ""},
    ],
    "true_false": [
        {"statement": "Animal cells have chloroplasts.", "answer": False},
        {"statement": "The nucleus stores DNA.", "answer": True},
    ],
    "match_column": [
        {"left": "Nucleus", "right": "Controls the cell"},
        {"left": "Chloroplast", "right": "Captures sunlight"},
        {"left": "Vacuole", "right": "Stores water"},
    ],
    "short_answer": [{"q": "Q1) Why are plant cells rigid?", "answer": "The cell wall.", "work_space_lines": 1}],
}


def _sections(doc):
    """Ordered (heading text, [non-empty paragraph texts under it]) pairs."""
    out = []
    for p in doc.paragraphs:
        if (p.style.name or "").startswith(("Heading", "Title")):
            out.append((p.text, []))
        elif out and p.text.strip():
            out[-1][1].append(p.text)
    return out


def test_numbering_restarts_per_section_and_key_matches_paper(tmp_path):
    out = worksheet.build(
        book={"grade": "Grade 7", "subject": "Science"},
        chapter={"title": "Cells", "sections": [{"section_title": "Cells", "content": "A cell is the basic unit."}]},
        analysis={},
        client=_StubClient(_NUMBERING_PAYLOAD),
        params={"num_questions": 10, "include_answer_key": True},
        out_dir=tmp_path,
    )
    doc = Document(str(out))

    # (c) Word's document-wide "List Number" counter must be gone entirely.
    assert all(p.style.name != "List Number" for p in doc.paragraphs)

    secs = _sections(doc)
    names = [name for name, _ in secs]
    key_at = names.index("Answer Key")
    paper = dict(secs[:key_at])
    key = dict(secs[key_at + 1:])

    # (a) every section restarts at 1 — True/False must not continue fill's count.
    assert paper["Section B — True or False"][0].startswith("1. ")

    # (d) an LLM-emitted "1." / "Q1)" prefix is stripped, not doubled.
    assert paper["Section A — Fill in the blanks"][0] == "1. The ____ controls the cell."
    assert paper["Section D — Short answer"][0] == "1. Why are plant cells rigid?"

    # (b) the key carries EXACTLY the paper's numbers, section by section.
    for name in ("Section A — Fill in the blanks", "Section B — True or False"):
        paper_nums = [t.split(".")[0] for t in paper[name]]
        key_nums = [t.split(".")[0] for t in key[name]]
        assert key_nums == paper_nums, name
    assert key["Section B — True or False"] == ["1. False", "2. True"]
    assert key["Section D — Short answer"] == ["1. The cell wall."]

    # A blank answer keeps its slot (em-dash), so numbering stays aligned.
    assert key["Section A — Fill in the blanks"] == ["1. nucleus", "2. —"]

    # The match key is "N. <letter>" (not the old self-numbered "1. 1 → C"),
    # and each letter points at the correct Column B row of the paper's table.
    match_key = key["Section C — Match the columns"]
    assert all(re.fullmatch(r"\d+\. [A-Z]", t) for t in match_key)
    right_by_letter = {}
    for row in doc.tables[0].rows[1:]:
        letter, right = row.cells[1].text.split(". ", 1)
        right_by_letter[letter] = right
    for i, line in enumerate(match_key):
        n, letter = line.split(". ")
        assert int(n) == i + 1
        assert right_by_letter[letter] == _NUMBERING_PAYLOAD["match_column"][i]["right"]


def test_write_worksheet_orders_kinds_and_keeps_match_whole(tmp_path):
    questions.write_worksheet(
        tmp_path, "T", "i",
        fill=[{"q": "a ____", "answer": "x"}],
        tf=[{"statement": "s", "answer": True}],
        match=[{"left": "L1", "right": "R1"}, {"left": "L2", "right": "R2"}],
        short=[{"q": "why?", "answer": "because"}],
    )
    payload = json.loads((tmp_path / "questions.json").read_text(encoding="utf-8"))
    assert [q["type"] for q in payload["questions"]] == ["fill_blank", "true_false", "match", "short"]
    assert payload["questions"][1]["answer"] is True  # true_false coerced to bool
